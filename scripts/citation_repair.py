#!/usr/bin/env python3
"""citation_repair.py - cite-formatter core logic as callable module.

Used by paper-check when detecting a references section.
Applies CNU《外国文学评论》citation rules with optional Better BibTeX .bib enrichment.
"""

from __future__ import annotations
import re
import sys
from pathlib import Path
from typing import Optional

# Allow running as script: ensure repo root is importable
_THIS = Path(__file__).resolve()
_REPO_ROOT = _THIS.parents[1]
if str(_REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(_REPO_ROOT))

from scripts.extract_references import detect_references_section, _is_list_item_stub


# -------------------------------------------------------------------
# Better BibTeX .bib file lookup
# -------------------------------------------------------------------

from scripts.lib.bib_parser import (
    parse_bib_file as _load_bib_file,
    build_title_index as _build_bib_index,
    search_by_title as _bib_search,
    merge_bib_into_parsed as _merge_bib_data,
    extract_pdf_path as _extract_pdf_path,
)

_BIB_ENTRIES: list[dict] | None = None
_BIB_TITLE_INDEX: dict[str, list[dict]] | None = None


# -------------------------------------------------------------------
# Citation parsing helpers
# -------------------------------------------------------------------

_RE_YEAR = re.compile(r"(?<!\d)((?:19|20)\d{2})(?!\d)")
_RE_STRIP_HTML = re.compile(r"<[^>]+>")
_RE_CITATION_TYPE = re.compile(
    r"\b(M|J|D|C|P|R|S|N|EB/OL|DSK|LED)\b",
    re.IGNORECASE,
)
_RE_EDITOR = re.compile(r"主编|编|ed\.|eds\.", re.IGNORECASE)
_RE_TRANSLATOR = re.compile(r"译|trans\.", re.IGNORECASE)
_RE_COMMA_CHINESE = re.compile(r"，")
_RE_PERIOD = re.compile(r"\.+")


class ParsedCitation(dict):
    """Structured citation data — keys: author, title, year, publisher,
    place, pages, type, language, translator, editor, raw.
    """

    pass


_RE_NUMBERED_PREFIX = re.compile(r"^\[?\d+\]?\s*")
_RE_CITATION_INDEX = re.compile(r"^\[(\d+)\]\s*")


def _apply_citation_fonts(para) -> None:
    """Apply 宋体 (CJK) + Times New Roman (Latin) at run level, 12pt."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt

    for run in para.runs:
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(12)


def parse_raw_citation(text: str) -> ParsedCitation:
    """
    Extract structural fields from a raw citation string.

    Handles both Chinese and English citation styles, including [N]-prefixed
    GB/T 7714 entries:
    - Chinese: [4] 作者：《书名》，出版地：出版社，年份，第X页。
    - English: [1] Author, "Title," *Journal*, vol. X, no. X (Year), pp. X-X.

    Returns a ParsedCitation dict with available fields filled.
    """
    original = text.strip()
    result: ParsedCitation = {"raw": original}

    # Extract and store citation index [N] before stripping
    idx_m = _RE_CITATION_INDEX.match(original)
    if idx_m:
        result["citation_index"] = int(idx_m.group(1))

    # Strip leading [N] or N. citation index
    normalized = _RE_NUMBERED_PREFIX.sub("", original).strip()

    # Detect language heuristic
    has_chinese = bool(re.search(r"[一-鿿]", normalized))
    result["language"] = "zh" if has_chinese else "en"

    # Strip HTML if present
    clean = _RE_STRIP_HTML.sub("", normalized)
    result["raw"] = clean
    normalized = clean

    # Year extraction
    year_match = _RE_YEAR.search(normalized)
    if year_match:
        result["year"] = year_match.group(1)

    # ---- Chinese-style citation parsing ----
    if has_chinese:
        # Author: text before the first 《
        author_match = re.match(r"^(.+?)[：:]?《", normalized)
        if author_match:
            result["author"] = author_match.group(1).rstrip("：:，, ").strip()

        # Title: text between first 《 》
        title_match = re.search(r"《(.+?)》", normalized)
        if title_match:
            result["title"] = title_match.group(1).strip()

        # Translator: word(s) between 》， and 译[，]  (e.g. 朱光潜译，)
        trans_match = re.search(r"》，([^，《》]+?)译[，,。]", normalized)
        if trans_match:
            result["translator"] = trans_match.group(1).strip()

        # Editor: text before 主编
        edit_match = re.search(r"([^，,]+?)主编", normalized)
        if edit_match:
            result["editor"] = edit_match.group(1).strip()

        # Place and publisher: look for 地：社 pattern after title
        # e.g. 北京：人民文学出版社 or 上海：上海译文出版社
        pub_match = re.search(r"》[^，]*，(?:[^，]+译，)?([^，]+)[：:]([^，，]+?)，(?:\d{4}|第)", normalized)
        if pub_match:
            result["place"] = pub_match.group(1).strip()
            result["publisher"] = pub_match.group(2).strip()
        elif not result.get("place"):
            # Simpler: find XX：XX between any comma segments
            for seg in re.split(r"[，,]", normalized):
                seg = seg.strip()
                m = re.match(r"^([^：:《》]{1,10})[：:]([^：:《》]{2,20})$", seg)
                if m and not re.search(r"[《》]", seg):
                    result["place"] = m.group(1)
                    result["publisher"] = m.group(2)
                    break

        # Year
        # Use non-digit boundaries to avoid Python 3 \b+CJK issue
        year_match = _RE_YEAR.search(normalized)
        if year_match:
            result["year"] = year_match.group(1)

        # Document type identifier [M], [J], etc. (GB/T 7714)
        type_match = _RE_CITATION_TYPE.search(normalized)
        if type_match:
            result["type"] = type_match.group(1).upper()

        # Page range
        page_match = re.search(r"第(\d+)[—\-~](\d+)页|第(\d+)页", normalized)
        if page_match:
            if page_match.group(1) and page_match.group(2):
                result["pages"] = f"{page_match.group(1)}—{page_match.group(2)}"
            elif page_match.group(3):
                result["pages"] = page_match.group(3)

        return result

    # ---- English-style citation parsing ----
    # Case 1: journal article — title in double quotes (ASCII or Unicode curly)
    quoted_title = re.search(r'[“"]([^”"]+)[”"]', normalized)
    if quoted_title:
        result["type"] = "J"
        result["title"] = quoted_title.group(1).rstrip(",").strip()
        # Author: everything before the opening quote, strip trailing comma/space
        result["author"] = normalized[: quoted_title.start()].rstrip(", ").strip()
        # Journal: text after closing ", " up to next comma or end
        # Reject matches that look like page references (p. X / pp. X-X)
        after = normalized[quoted_title.end():].lstrip(", ")
        journal_m = re.match(r"\*?([^,*]+?)\*?(?:,|$)", after)
        if journal_m:
            j = journal_m.group(1).strip()
            if not re.match(r"^pp?\.\s*\d", j, re.IGNORECASE):
                result["journal"] = j
        # vol / no / year / pages
        for pat, key in [
            (r"vol\.\s*(\d+)", "volume"),
            (r"no\.\s*(\d+)", "no"),
            (r"pp?\.\s*([\d–\-]+)", "pages"),
        ]:
            m = re.search(pat, normalized, re.IGNORECASE)
            if m:
                result[key] = m.group(1)
        return result

    # Case 2: book with *Title* italics marker
    italic_title = re.search(r"\*([^*]+)\*", normalized)
    if italic_title:
        result["title"] = italic_title.group(1).strip()
        result["author"] = normalized[: italic_title.start()].rstrip(", ").strip()
        after = normalized[italic_title.end():].lstrip(", ")
        pub_m = re.match(r"(.+?:\s*.+?),\s*(\d{4})", after)
        if pub_m:
            pub_info = pub_m.group(1)
            if ":" in pub_info:
                p = pub_info.split(":", 1)
                result["place"] = p[0].strip()
                result["publisher"] = p[1].strip()
        return result

    # Case 3: GB/T-style edited volume — Title, Editor(s) eds. Place: Publisher, Year.
    eds_m = re.search(r",\s+(.+?)\s+eds?\.\s+(.+?:\s*.+?),\s*(\d{4})", normalized, re.IGNORECASE)
    if eds_m:
        result["title"] = normalized.split(",")[0].strip()
        result["editor"] = eds_m.group(1).strip()
        pub_info = eds_m.group(2)
        if ":" in pub_info:
            p = pub_info.split(":", 1)
            result["place"] = p[0].strip()
            result["publisher"] = p[1].strip()
        return result

    # Fallback: split on comma, grab first segment as author
    result["author"] = normalized.split(",")[0].strip() if "," in normalized else normalized
    return result


# -------------------------------------------------------------------
# CNU format rules (from cite-formatter/references/format-rules.md)
# -------------------------------------------------------------------

def format_cnu_citation(parsed: ParsedCitation) -> str:
    """
    Format a parsed citation according to CNU《外国文学评论》rules.

    Rules by language and type:
    - Chinese book:        作者：《书名》，出版地：出版社，年份，第X页。
    - Chinese journal:     作者：《文章题名》，《期刊名》年份第X期。
    - Chinese chapter:     作者：《章节名》，编者编：《书名》，出版地：出版社，年份，第X—X页。
    - English book:        Author, *Title*, Place: Publisher, Year, p. X.
    - English journal:     Author, "Article," *Journal*, vol. X, no. X (Year), pp. X-X.
    - English chapter:     Author, "Chapter," in Editor, ed., *Book*, Place: Publisher, Year, p. X.
    - Subsequent (repeat):Author, *Short Title*, p. X.
    """
    lang = parsed.get("language", "zh")
    ctype = parsed.get("type", "M")

    if lang == "zh":
        return _format_cnu_chinese(parsed)
    else:
        return _format_cnu_english(parsed, ctype)


def _format_cnu_chinese(p: ParsedCitation) -> str:
    """Format a Chinese citation per CNU rules."""
    author = p.get("author", "")
    title = p.get("title", "")
    place = p.get("place", "")
    publisher = p.get("publisher", "")
    year = p.get("year", "")
    pages = p.get("pages", "")
    translator = p.get("translator", "")
    editor = p.get("editor", "")

    # Determine format based on type
    ctype = p.get("type", "M")

    if ctype == "J":
        # Journal article
        journal = p.get("journal", p.get("publisher", ""))
        issue = p.get("issue", "")
        formatted = f"{author}：《{title}》，《{journal}》{year}"
        if issue:
            formatted += f"第{issue}期"
        formatted += "。"
        return formatted

    elif ctype in ("C", "M"):  # Book or chapter
        parts = [f"{author}：《{title}》"]
        if translator:
            parts.append(f"{translator}译")
        if place and publisher:
            parts.append(f"{place}：{publisher}")
        elif publisher:
            parts.append(publisher)
        if year:
            parts.append(f"{year}年")
        if pages:
            parts.append(f"第{pages}页")
        # Build result string avoiding double commas
        result = parts[0]
        for part in parts[1:]:
            if result.endswith("，") or result.endswith("："):
                result += part
            else:
                result += "，" + part
        return result + "。"

    else:
        # Default: book format
        parts = [f"{author}：《{title}》"]
        if place and publisher:
            parts.append(f"{place}：{publisher}")
        elif publisher:
            parts.append(publisher)
        if year:
            parts.append(f"{year}年")
        if pages:
            parts.append(f"，第{pages}页")
        result = "，".join(parts) + "。"
        return result


def _format_cnu_english(p: ParsedCitation, ctype: str = "M") -> str:
    """Format an English citation per CNU rules."""
    author = p.get("author", "")
    title = p.get("title", "")
    place = p.get("place", "")
    publisher = p.get("publisher", "")
    year = p.get("year", "")
    pages = p.get("pages", "")
    journal = p.get("journal", "")
    volume = p.get("volume", "")
    issue = p.get("issue", "")
    no = p.get("no", "")

    if ctype == "J":
        if journal:
            # Full journal article: Author, "Title," *Journal*, vol. X, no. X (Year), pp. X-X.
            vol_no_parts = []
            if volume:
                vol_no_parts.append(f"vol. {volume}")
            if no:
                vol_no_parts.append(f"no. {no}")
            vol_no_str = ", ".join(vol_no_parts)
            if vol_no_str and year:
                vol_no_str += f" ({year})"
            elif year:
                vol_no_str = f"({year})"
            parts = [f"{author}, \"{title},\" *{journal}*"]
            if vol_no_str:
                parts.append(vol_no_str)
            if pages:
                pages_norm = re.sub(r"\s+", "", pages)
                parts.append(f"pp. {pages_norm}")
            return ", ".join(parts) + "."
        else:
            # Subsequent citation — Author, "Title," p. X.
            title_part = f'"{title},"' if title else ""
            parts = [f"{author}, {title_part}"] if author else [title_part]
            if pages:
                pages_norm = re.sub(r"\s+", "", pages)
                prefix = "pp." if re.search(r"\d[-–]\d", pages_norm) else "p."
                parts.append(f"{prefix} {pages_norm}")
            return " ".join(parts).rstrip(", ") + "."

    elif ctype == "C" and place and publisher:
        # Book chapter
        return (
            f"{author}, \"{title},\" in *{p.get('book_title', title)}*, "
            f"ed. by {p.get('editor', author)}, "
            f"{place}: {publisher}, {year}, p. {pages}."
        )

    else:
        # Book (or edited volume)
        if not author and p.get("editor"):
            author = p["editor"] + ", eds."
        parts = [f"{author}, *{title}*" if author else f"*{title}*"]
        if place and publisher:
            parts.append(f"{place}: {publisher}")
        elif publisher:
            parts.append(publisher)
        if year:
            parts.append(str(year))
        if pages:
            parts.append(f"p. {pages}" if not pages.startswith("pp.") else pages)
        return ", ".join(parts) + "."


# -------------------------------------------------------------------
# P1/P2/P3 severity classification
# -------------------------------------------------------------------

class CitationWarning:
    def __init__(self, severity: str, message: str):
        self.severity = severity  # P1, P2, P3
        self.message = message

    def __repr__(self):
        return f"⚠️ [{self.severity}] {self.message}"


def classify_warnings(parsed: ParsedCitation) -> list[CitationWarning]:
    """
    Classify missing or incorrect fields by severity.

    P1 (must fix): Missing required fields (author/title/year/publisher);
                   book title format wrong (Chinese should use《》, English *italic*)
    P2 (should fix): Page range wrong (should use — not -); punctuation errors
    P3 (suggested):  Subsequent citation not simplified; citekey spelling risk
    """
    warnings: list[CitationWarning] = []
    lang = parsed.get("language", "zh")
    title = parsed.get("title", "")

    # P1: Missing required fields
    if not parsed.get("author") and not parsed.get("editor"):
        warnings.append(CitationWarning("P1", "缺少必填字段：作者"))
    if not title:
        warnings.append(CitationWarning("P1", "缺少必填字段：题名"))
    if not parsed.get("year") and not parsed.get("is_subsequent"):
        warnings.append(CitationWarning("P1", "缺少必填字段：年份"))

    # P1: Book title format — check ORIGINAL text, not extracted title field
    raw = parsed.get("raw", "")
    if title:
        if lang == "zh":
            # Warn only if original text lacks 《》 markers entirely
            if "《" not in raw and "》" not in raw:
                warnings.append(
                    CitationWarning("P1", "中文书名应使用《》")
                )
        else:
            # For books only (not journal articles): check for italic markers
            ctype_check = parsed.get("type", "M")
            if ctype_check not in ("J",) and "*" not in raw and "_" not in raw:
                warnings.append(
                    CitationWarning("P2", "英文书名建议使用斜体（*Title*）")
                )

    # P1: Missing publisher/place for books (not journals, not subsequent citations)
    ctype = parsed.get("type", "M")
    if ctype not in ("J",) and not parsed.get("is_subsequent"):
        if not parsed.get("publisher") and not parsed.get("editor"):
            warnings.append(CitationWarning("P1", "缺少出版社信息"))

    # P2: Page range punctuation
    pages = parsed.get("pages", "")
    if pages:
        if "－" in pages or (re.search(r"\d+-\d+", pages) and lang == "zh"):
            warnings.append(
                CitationWarning("P2", "中文页码范围应用全角破折号（—）")
            )
        # Fire only when raw text uses singular "p." for a range (not "pp.")
        if lang == "en":
            raw_page_ref = re.search(r"\bp\.(?: )?\d+\s*[-–]\s*\d+", parsed.get("raw", ""), re.IGNORECASE)
            if raw_page_ref and not raw_page_ref.group().lower().startswith("pp"):
                warnings.append(
                    CitationWarning("P2", "英文页码应使用 pp. 格式")
                )

    # P2: Chinese comma instead of colon for place:publisher
    raw = parsed.get("raw", "")
    if "：" in raw or "，" in raw:
        # Check if format looks like place：publisher
        if re.search(r"[^（）\s]+[：:][^，,]+[，,]\d{4}", raw):
            pass  # Correct format
        elif re.search(r"\S[，,]\S.*\d{4}", raw):
            warnings.append(
                CitationWarning("P2", "检查格式：出版地与出版社间应用冒号分隔")
            )

    # Note: is_subsequent=True means citation is already in simplified re-citation form
    # (correct CNU behavior) — no warning needed here.

    return warnings


# -------------------------------------------------------------------
# Main entry point: process a list of raw citations
# -------------------------------------------------------------------

class CitationRepairResult(dict):
    """Result for one citation: original, formatted, parsed, warnings,
    bib_status, needs_manual_review."""

    pass


def process_citations(
    raw_citations: list[str],
    bib_path: str | Path | None = None,
    spec_preference: str = "CNU",
) -> list[CitationRepairResult]:
    """
    Process a list of raw citation strings.

    Args:
        raw_citations: list of citation paragraph texts
        bib_path: path to Better BibTeX .bib file for metadata enrichment
        spec_preference: "CNU" (default) or "GB/T7714"

    Returns:
        list of CitationRepairResult dicts with keys:
          original, formatted, parsed_fields, warnings (list of strings),
          bib_status, needs_manual_review
    """
    global _BIB_ENTRIES, _BIB_TITLE_INDEX

    # Load .bib file if provided
    if bib_path:
        _BIB_ENTRIES = _load_bib_file(bib_path)
        _BIB_TITLE_INDEX = _build_bib_index(_BIB_ENTRIES)
        print(f"[citation_repair] Loaded {len(_BIB_ENTRIES)} entries from {bib_path}")
    else:
        _BIB_ENTRIES = None
        _BIB_TITLE_INDEX = None

    results: list[CitationRepairResult] = []

    for citation_text in raw_citations:
        result = _process_single(
            citation_text,
            spec_preference=spec_preference,
        )
        results.append(result)

    return results


def _process_single(
    citation_text: str,
    spec_preference: str = "CNU",
    parsed_overrides: dict | None = None,
) -> CitationRepairResult:
    """Process one citation."""
    original = citation_text.strip()
    result = CitationRepairResult(original=original)

    # 1. Parse fields from raw text
    parsed = parse_raw_citation(original)
    if parsed_overrides:
        parsed.update(parsed_overrides)
    result["parsed_fields"] = dict(parsed)

    # 2. Better BibTeX .bib lookup if available
    bib_status = "not_checked"
    if _BIB_ENTRIES and _BIB_TITLE_INDEX:
        bib_matches = _bib_search(parsed, _BIB_ENTRIES, _BIB_TITLE_INDEX)
        if len(bib_matches) == 1:
            parsed = _merge_bib_data(parsed, bib_matches[0])
            bib_status = "matched"
        elif len(bib_matches) > 1:
            # Multiple candidates — pick best by year match
            if parsed.get("year"):
                year_matches = [b for b in bib_matches if str(b.get("year_int", "")) == parsed["year"]]
                if len(year_matches) == 1:
                    parsed = _merge_bib_data(parsed, year_matches[0])
                    bib_status = "matched"
                else:
                    bib_status = f"multiple_candidates({len(bib_matches)})"
            else:
                bib_status = f"multiple_candidates({len(bib_matches)})"
        else:
            bib_status = "not_found"

    result["bib_status"] = bib_status
    result["pdf_path"] = parsed.get("pdf_path", "")

    # 3. Format with CNU rules
    formatted = format_cnu_citation(parsed)
    result["formatted"] = formatted

    # 4. Classify warnings
    warnings = classify_warnings(parsed)
    result["warnings"] = [str(w) for w in warnings]
    result["needs_manual_review"] = bib_status in (
        "multiple_candidates",
        "not_found",
    ) or any(w.severity == "P1" for w in warnings)

    return result


# -------------------------------------------------------------------
# Write corrected citations back to .docx
# -------------------------------------------------------------------

def repair_citations_in_docx(
    docx_path: str | Path,
    citation_results: list[CitationRepairResult],
    output_path: str | Path | None = None,
    doc=None,
) -> None:
    """
    Replace original citation paragraphs with formatted versions in .docx.
    Skips citations marked needs_manual_review.

    Pass an already-open Document via `doc` to avoid a redundant open/save
    cycle (e.g. when called from fix_format which manages the document itself).
    When `doc` is provided, no file is written — the caller is responsible for
    saving.
    """
    from docx import Document

    _own_doc = doc is None
    if _own_doc:
        doc = Document(docx_path)
    out_path = Path(output_path) if output_path else Path(docx_path)

    ref_result = detect_references_section(docx_path)
    if not ref_result.found:
        return

    citation_map = {
        i: r
        for i, r in enumerate(ref_result.raw_citations)
        if not citation_results[i].get("needs_manual_review", False)
    }

    para_idx = ref_result.section_start or 0
    replacement_count = 0

    for para in doc.paragraphs[para_idx:]:
        if ref_result.section_end and para_idx >= ref_result.section_end:
            break
        text = para.text.strip()
        if not text or _is_list_item_stub(text):
            para_idx += 1
            continue
        for i, raw_text in citation_map.items():
            if raw_text == text:
                formatted = citation_results[i].get("formatted", "")
                if formatted:
                    # Restore [N] prefix from original
                    idx_m = _RE_CITATION_INDEX.match(raw_text)
                    if idx_m:
                        formatted = f"[{idx_m.group(1)}] {formatted}"
                    para.text = formatted
                    _apply_citation_fonts(para)
                    replacement_count += 1
                break
        para_idx += 1

    if _own_doc:
        doc.save(out_path)


# -------------------------------------------------------------------
# Standalone report generation
# -------------------------------------------------------------------

def generate_citation_report(
    citation_results: list[CitationRepairResult],
    output_path: str | Path,
    spec_preference: str = "CNU",
) -> None:
    """
    Generate a standalone reference format report .docx.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading("参考文献格式报告", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"格式规范：{spec_preference}（CNU《外国文学评论》2024修订版）")
    doc.add_paragraph(f"共检测 {len(citation_results)} 条引用")

    total_p1 = sum(
        1 for r in citation_results
        if any("P1" in w for w in r.get("warnings", []))
    )
    total_p2 = sum(
        1 for r in citation_results
        if any("P2" in w for w in r.get("warnings", []))
    )
    total_p3 = sum(
        1 for r in citation_results
        if any("P3" in w for w in r.get("warnings", []))
    )
    doc.add_paragraph(
        f"问题分布：P1(必须修复) {total_p1} 条，P2(应当修复) {total_p2} 条，P3(建议修复) {total_p3} 条"
    )

    doc.add_paragraph("")

    # Per-citation entries
    for idx, r in enumerate(citation_results, 1):
        p = doc.add_paragraph()
        p.add_run(f"[{idx}] ").bold = True
        p.add_run(r["original"][:100])
        if len(r["original"]) > 100:
            p.add_run("...")

        formatted = r.get("formatted", "")
        if formatted:
            pf = doc.add_paragraph()
            pf.add_run("  规范化：").bold = True
            pf.add_run(formatted)

        warnings = r.get("warnings", [])
        if warnings:
            for w in warnings:
                pw = doc.add_paragraph()
                pw.add_run(f"  {w}").bold = False

        bib = r.get("bib_status", r.get("zotero_status", ""))
        if bib and bib != "not_checked":
            doc.add_paragraph(f"  BibTeX状态：{bib}")

        manual = r.get("needs_manual_review", False)
        if manual:
            pm = doc.add_paragraph()
            pm.add_run("  ⚠️ 需要人工核实").bold = True

        doc.add_paragraph("")

    doc.save(output_path)


def process_footnote_citations(
    footnote_items,
    bib_path: str | Path | None = None,
) -> list[CitationRepairResult]:
    """
    Process footnotes from extract_footnotes.FootnoteItem list.

    - Skips prose-only footnotes.
    - Marks subsequent citations so publisher/place warnings are suppressed.
    - Uses cite_text (prose prefix stripped) for parsing, preserves raw_text.
    """
    results: list[CitationRepairResult] = []
    for item in footnote_items:
        if item.is_prose_only:
            continue
        overrides = {"is_subsequent": True} if item.is_subsequent else None
        result = _process_single(
            item.cite_text,
            parsed_overrides=overrides,
        )
        result["footnote_id"] = item.footnote_id
        result["raw_footnote_text"] = item.raw_text  # full text incl. prose prefix
        results.append(result)
    return results


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Citation repair CLI")
    parser.add_argument("thesis", help="Path to thesis .docx")
    parser.add_argument("--output", "-o", help="Output path for report .docx")
    parser.add_argument(
        "--bib",
        help="Path to Better BibTeX .bib file for metadata enrichment",
    )
    args = parser.parse_args()

    result = detect_references_section(args.thesis)
    if not result.found:
        print("No references section found.")
        sys.exit(0)

    print(f"Found {result.citation_count} citations. Processing...")
    processed = process_citations(
        result.raw_citations,
        bib_path=args.bib,
    )

    for i, r in enumerate(processed, 1):
        print(f"\n[{i}] {r['original'][:60]}...")
        print(f"  → {r['formatted']}")
        if r["warnings"]:
            for w in r["warnings"]:
                print(f"  {w}")

    if args.output:
        generate_citation_report(processed, args.output)
        print(f"\nReport saved to {args.output}")