#!/usr/bin/env python3
"""check_fake_layout.py - 检测并清理"假排版"。

"假排版"指用排版无关的手段冒充版式的做法，常见于学生稿：
  - 连续空段落堆叠制造垂直间距（应由段前/段后间距控制）
  - 段首手工空格冒充首行缩进（应由 first_line_indent 控制）
  - 段内手工换行(w:br)冒充分段或行距
  - 行尾游离空格
这些会让样式系统失效、跨设备错位，必须在统一样式前清掉。
"""
from __future__ import annotations
import argparse
import json
import sys
from pathlib import Path

_THIS = Path(__file__).resolve()
if str(_THIS.parents[1]) not in sys.path:
    sys.path.insert(0, str(_THIS.parents[1]))

from docx import Document
from docx.oxml.ns import qn

_LEADING_SPACE = "　 \t"  # 全角空格 / 半角空格 / 制表符（单字符成员判断 + lstrip 两用）


def _is_blank(p) -> bool:
    return not p.text.strip()


def detect_fake_layout(docx_path: str) -> dict:
    """Return {file, total, issues:[{type, location, detail, count?, autofix_safe}]}."""
    doc = Document(docx_path)
    paras = doc.paragraphs
    issues = []

    # 1) 连续空段（>=2）
    run_len = 0
    run_start = 0
    for i, p in enumerate(paras + [None]):
        if p is not None and _is_blank(p):
            if run_len == 0:
                run_start = i
            run_len += 1
        else:
            if run_len >= 2:
                issues.append({
                    "type": "consecutive_blank", "location": f"段 {run_start}-{run_start + run_len - 1}",
                    "count": run_len, "detail": f"{run_len} 个连续空段，疑似手工制造间距",
                    "autofix_safe": True,
                })
            run_len = 0

    # 2) 段首手工空格 / 3) 段内手工换行 / 4) 行尾空格
    for i, p in enumerate(paras):
        t = p.text
        if not t.strip():
            continue
        if t[:1] in _LEADING_SPACE:
            n = len(t) - len(t.lstrip(_LEADING_SPACE))
            issues.append({
                "type": "manual_indent_spaces", "location": f"段 {i}",
                "detail": f"段首 {n} 个手工空格，疑似冒充首行缩进/居中：{t[:20]!r}",
                "autofix_safe": True,
            })
        if len(p._p.findall(".//" + qn("w:br"))) > 0:
            issues.append({
                "type": "intra_para_break", "location": f"段 {i}",
                "detail": f"段内含手工换行(w:br)：{t[:20]!r}",
                "autofix_safe": False,
            })
        if t != t.rstrip() and t.strip():
            issues.append({
                "type": "trailing_space", "location": f"段 {i}",
                "detail": "行尾游离空格", "autofix_safe": True,
            })

    return {"file": docx_path, "total": len(issues), "issues": issues}


def clean_fake_layout(doc) -> int:
    """In-place 清理可安全自动处理的假排版，返回处理动作数。

    只做高确定性的清理：折叠连续空段为单个、去段首手工空格、去行尾空格。
    段内手工换行(w:br)不自动删（可能有意），仅由 detect 标记待人工确认。
    """
    actions = 0

    # 1) 折叠连续空段：保留每段连续空白的第 1 个，删除其余
    prev_blank = False
    for p in list(doc.paragraphs):
        blank = not p.text.strip()
        if blank and prev_blank:
            p._p.getparent().remove(p._p)
            actions += 1
        prev_blank = blank

    # 2) 段首手工空格 + 3) 行尾空格
    # 段首/段尾的手工空格可能跨多个 run（Word 常把空格单独拆成一个 run）。
    # 只改 runs[0]/runs[-1] 会漏掉这种情况，且使动作计数与实际不符。
    # 因此从两端逐个 run 推进：整段为空白的 run 清空并继续，遇到含实义内容的
    # run 则 strip 其外侧空白后停止。每段最多计 1 次首部 + 1 次尾部动作。
    for p in doc.paragraphs:
        if not p.text.strip() or not p.runs:
            continue
        # 段首
        if p.text[:1] in _LEADING_SPACE:
            for run in p.runs:
                if run.text == "":
                    continue
                new = run.text.lstrip(_LEADING_SPACE)
                run.text = new
                if new != "":  # 该 run 含实义内容，段首空白处理完毕
                    break
            actions += 1
        # 段尾
        if p.text[-1:] in _LEADING_SPACE:
            for run in reversed(p.runs):
                if run.text == "":
                    continue
                new = run.text.rstrip(_LEADING_SPACE)
                run.text = new
                if new != "":
                    break
            actions += 1

    return actions


def main() -> None:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")
    ap = argparse.ArgumentParser(description="检测假排版（空行堆叠/手工空格/段内换行/行尾空格）")
    ap.add_argument("docx")
    ap.add_argument("--output", "-o")
    args = ap.parse_args()
    res = detect_fake_layout(args.docx)
    if args.output:
        Path(args.output).write_text(json.dumps(res, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"检测到 {res['total']} 处假排版 → {args.output}")
    else:
        print(f"检测到 {res['total']} 处假排版：")
        for it in res["issues"]:
            print(f"  [{it['type']}] {it['location']} {it['detail']}")


if __name__ == "__main__":
    main()
