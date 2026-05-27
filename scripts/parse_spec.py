#!/usr/bin/env python3
"""parse_spec.py - Extract formatting rules from a sample .docx document.

Given a correctly-formatted thesis/journal paper, extract its formatting
specifications (margins, fonts, heading styles, table styles, headers/footers)
and output a structured JSON spec file.

This spec file can then be used by check_format.py and fix_format.py as the
target formatting standard, replacing hardcoded defaults.

Usage:
    python parse_spec.py sample.docx --output spec.json
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any


def _pt(length) -> float | None:
    """Convert a python-docx Length to pt, or return None."""
    try:
        return length.pt if length is not None else None
    except Exception:
        return None


def _mm(length) -> float | None:
    """Convert EMU to mm."""
    try:
        return length / 914400 * 25.4 if length is not None else None
    except Exception:
        return None


def _get_font_info(run) -> dict:
    """Extract font info from a run, including CJK eastAsia font via XML."""
    from docx.oxml.ns import qn

    info: dict[str, Any] = {}
    rPr = run._r.get_or_add_rPr() if hasattr(run, '_r') else None
    if rPr is not None:
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is not None:
            info["font_latin"] = rFonts.get(qn("w:ascii"), "")
            info["font_cjk"] = rFonts.get(qn("w:eastAsia"), "")
            info["font_hAnsi"] = rFonts.get(qn("w:hAnsi"), "")
    if run.font.size:
        info["size"] = _pt(run.font.size)
    if run.font.bold:
        info["bold"] = True
    if run.font.italic:
        info["italic"] = True
    return info


def parse_page_setup(doc) -> dict:
    """Extract page setup from the first section."""
    setup: dict[str, Any] = {}
    try:
        section = doc.sections[0]
        pw = section.page_width
        ph = section.page_height
        if pw is not None and ph is not None:
            w_mm = _mm(pw)
            h_mm = _mm(ph)
            setup["paper"] = f"{w_mm:.0f}×{h_mm:.0f}mm"
            if w_mm and abs(w_mm - 210) < 5 and abs(h_mm - 297) < 5:
                setup["paper"] = "A4"
            elif w_mm and abs(w_mm - 210) < 5 and abs(h_mm - 297) < 5:
                setup["paper"] = "A4"

        margins = {}
        for side in ("top", "bottom", "left", "right"):
            val = getattr(section, f"{side}_margin", None)
            mm_val = _mm(val)
            if mm_val is not None:
                margins[side] = round(mm_val, 1)
        if margins:
            setup["margins"] = margins

        # Header/footer distances
        hd = _mm(section.header_distance) if section.header_distance else None
        fd = _mm(section.footer_distance) if section.footer_distance else None
        if hd is not None:
            setup["header_distance"] = round(hd, 1)
        if fd is not None:
            setup["footer_distance"] = round(fd, 1)
    except (IndexError, AttributeError):
        pass
    return setup


def parse_styles(doc) -> dict:
    """Extract Normal and Heading 1-4 style definitions."""
    from docx.oxml.ns import qn

    styles: dict[str, Any] = {}
    target_names = ["Normal"] + [f"Heading {i}" for i in range(1, 5)]

    for style_name in target_names:
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue

        info: dict[str, Any] = {}

        # Font from style
        font = style.font
        rPr = font._element
        if rPr is not None:
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is not None:
                east_asia = rFonts.get(qn("w:eastAsia"), "")
                ascii_f = rFonts.get(qn("w:ascii"), "")
                if east_asia:
                    info["font_cjk"] = east_asia
                if ascii_f:
                    info["font_latin"] = ascii_f

        if font.size:
            info["size"] = _pt(font.size)
        if font.bold:
            info["bold"] = True

        # Paragraph format from style
        pf = style.paragraph_format
        if pf.alignment is not None:
            align_map = {0: "left", 1: "center", 2: "right", 3: "justify"}
            info["alignment"] = align_map.get(pf.alignment, str(pf.alignment))

        ls = pf.line_spacing
        if ls is not None:
            from docx.enum.text import WD_LINE_SPACING
            rule = pf.line_spacing_rule
            if rule is not None and rule != WD_LINE_SPACING.MULTIPLE:
                ratio = _pt(ls) / 12.0 if _pt(ls) else None
            else:
                ratio = float(ls) if isinstance(ls, (int, float)) else None
            if ratio is not None:
                info["line_spacing"] = round(ratio, 2)

        fi = pf.first_line_indent
        if fi is not None:
            from docx.shared import Pt
            # Convert EMU to character units (approximate at body size)
            body_size = info.get("size", 12)
            chars = _pt(fi) / body_size if _pt(fi) else None
            if chars is not None:
                info["first_line_indent_chars"] = round(chars, 1)

        sb = pf.space_before
        sa = pf.space_after
        if sb is not None:
            info["space_before"] = _pt(sb)
        if sa is not None:
            info["space_after"] = _pt(sa)

        if info:
            styles[style_name] = info

    return styles


def parse_paragraph_samples(doc, limit: int = 50) -> dict:
    """Sample body paragraphs for formatting patterns."""
    from docx.enum.text import WD_LINE_SPACING

    samples: dict[str, Any] = {
        "count": 0,
        "first_line_indent_chars": [],
        "line_spacing_values": [],
        "space_before_values": [],
        "space_after_values": [],
    }

    count = 0
    for para in doc.paragraphs:
        sname = para.style.name if para.style else ""
        if sname.startswith("Heading") or sname.startswith("TOC"):
            continue
        text = para.text.strip()
        if not text or len(text) < 10:
            continue

        count += 1
        if count > limit:
            break

        # First line indent
        fi = para.paragraph_format.first_line_indent
        if fi is not None:
            size_pt = 12  # default assumption
            try:
                if para.runs and para.runs[0].font.size:
                    size_pt = _pt(para.runs[0].font.size) or 12
            except Exception:
                pass
            chars = _pt(fi) / size_pt if _pt(fi) else None
            if chars is not None and 0.5 < chars < 5:
                samples["first_line_indent_chars"].append(round(chars, 1))

        # Line spacing
        ls = para.paragraph_format.line_spacing
        if ls is not None:
            rule = para.paragraph_format.line_spacing_rule
            try:
                if rule is not None and rule == WD_LINE_SPACING.MULTIPLE:
                    ratio = float(ls)
                else:
                    ratio = _pt(ls) / 12.0 if _pt(ls) else None
                if ratio is not None and 0.8 < ratio < 3.0:
                    samples["line_spacing_values"].append(round(ratio, 2))
            except Exception:
                pass

        # Space before/after
        sb = para.paragraph_format.space_before
        sa = para.paragraph_format.space_after
        if sb is not None:
            samples["space_before_values"].append(_pt(sb))
        if sa is not None:
            samples["space_after_values"].append(_pt(sa))

    samples["count"] = count
    return samples


def parse_tables(doc) -> list[dict]:
    """Analyze table border styles (three-line table detection)."""
    from docx.oxml.ns import qn

    tables_info: list[dict] = []
    for idx, table in enumerate(doc.tables):
        tbl_pr = table._tbl.tblPr
        if tbl_pr is None:
            tables_info.append({"index": idx, "style": "unknown"})
            continue

        borders_el = tbl_pr.find(qn("w:tblBorders"))
        border_info: dict[str, Any] = {"index": idx}

        if borders_el is not None:
            for side in ("top", "bottom", "left", "right", "insideH", "insideV"):
                side_el = borders_el.find(qn(f"w:{side}"))
                if side_el is not None:
                    val = side_el.get(qn("w:val"), "none")
                    sz = side_el.get(qn("w:sz"), "0")
                    border_info[f"border_{side}_val"] = val
                    border_info[f"border_{side}_sz"] = int(sz) if sz.isdigit() else 0

        # Check for three-line table pattern:
        # top + bottom thick, no left/right, insideH thin for header row only
        has_top = border_info.get("border_top_val", "none") not in ("none", "nil")
        has_bottom = border_info.get("border_bottom_val", "none") not in ("none", "nil")
        no_left = border_info.get("border_left_val", "none") in ("none", "nil")
        no_right = border_info.get("border_right_val", "none") in ("none", "nil")

        if has_top and has_bottom and no_left and no_right:
            border_info["style"] = "three_line"
        elif has_top and has_bottom:
            border_info["style"] = "boxed"
        else:
            border_info["style"] = "other"

        # Row count
        border_info["row_count"] = len(table.rows)
        border_info["col_count"] = len(table.columns)

        tables_info.append(border_info)

    return tables_info


def parse_headers_footers(doc) -> dict:
    """Extract header and footer content."""
    result: dict[str, Any] = {"sections": []}

    for idx, section in enumerate(doc.sections):
        sec_info: dict[str, Any] = {"index": idx}

        # Different first page header
        try:
            sec_info["different_first_page"] = bool(section.different_first_page_header_footer)
        except Exception:
            pass

        # Header
        try:
            header = section.header
            if header and header.paragraphs:
                texts = [p.text.strip() for p in header.paragraphs if p.text.strip()]
                sec_info["header_text"] = " ".join(texts)
                # Extract font info from first run
                for p in header.paragraphs:
                    if p.runs:
                        fi = _get_font_info(p.runs[0])
                        if fi:
                            sec_info["header_font"] = fi
                        break
        except Exception:
            pass

        # Footer / page numbers
        try:
            footer = section.footer
            if footer and footer.paragraphs:
                texts = [p.text.strip() for p in footer.paragraphs if p.text.strip()]
                sec_info["footer_text"] = " ".join(texts)
                # Check for page number field
                for p in footer.paragraphs:
                    for run in p.runs:
                        fld = run._r.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldChar")
                        if fld:
                            sec_info["has_page_number"] = True
                            break
        except Exception:
            pass

        result["sections"].append(sec_info)

    return result


def parse_spec(docx_path: str | Path) -> dict:
    """Extract formatting specifications from a sample .docx file.

    Returns a structured dict suitable for JSON serialization.
    """
    from docx import Document

    docx_path = Path(docx_path)
    if not docx_path.exists():
        return {"error": f"文件不存在: {docx_path}"}
    if docx_path.suffix.lower() != ".docx":
        return {"error": f"仅支持 .docx 格式: {docx_path.suffix}"}

    try:
        doc = Document(docx_path)
    except Exception as e:
        return {"error": f"无法打开文档: {e}"}

    spec: dict[str, Any] = {
        "source": str(docx_path),
        "page_setup": parse_page_setup(doc),
        "styles": parse_styles(doc),
        "paragraph_samples": parse_paragraph_samples(doc),
        "tables": parse_tables(doc),
        "headers_footers": parse_headers_footers(doc),
    }

    # Derive summary values from paragraph samples
    samples = spec["paragraph_samples"]
    if samples["first_line_indent_chars"]:
        from statistics import median
        spec["paragraphs_summary"] = {
            "first_line_indent_chars": round(
                median(samples["first_line_indent_chars"]), 1
            ),
            "line_spacing": round(median(samples["line_spacing_values"]), 2)
            if samples["line_spacing_values"]
            else None,
            "space_before_pt": round(median(samples["space_before_values"]), 1)
            if samples["space_before_values"]
            else None,
            "space_after_pt": round(median(samples["space_after_values"]), 1)
            if samples["space_after_values"]
            else None,
        }

    return spec


def main():
    import sys as _sys
    _p = Path(__file__).resolve().parents[1]
    if str(_p) not in _sys.path:
        _sys.path.insert(0, str(_p))
    from scripts.lib import setup_utf8_output
    setup_utf8_output()
    parser = argparse.ArgumentParser(
        description="Extract formatting rules from a sample .docx"
    )
    parser.add_argument("sample", help="Path to sample .docx")
    parser.add_argument("--output", "-o", help="Output JSON path")
    args = parser.parse_args()

    spec = parse_spec(args.sample)

    if args.output:
        Path(args.output).write_text(
            json.dumps(spec, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        print(f"Spec written to {args.output}")
    else:
        print(json.dumps(spec, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
