#!/usr/bin/env python3
"""generate_template.py - Generate pre-formatted Chinese academic paper templates.

Creates a .docx file with all styles, margins, and section structure
pre-configured according to Chinese academic standards (GB/T 7713.1-2006).

Modes:
  - journal (default): title page, abstract, keywords, body, references
  - thesis: cover, abstract, TOC, body, references, acknowledgments, appendices

Supports spec.json from parse_spec.py for custom formatting.

Usage:
    python generate_template.py --mode journal --output "论文模板.docx"
    python generate_template.py --mode thesis --output "毕业论文模板.docx"
    python generate_template.py --spec spec.json --output "自定义模板.docx"
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any

_THIS = Path(__file__).resolve()
_REPO_ROOT = _THIS.parents[1]
if str(_REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(_REPO_ROOT))


# -------------------------------------------------------------------
# Default template settings
# -------------------------------------------------------------------

DEFAULT_SETTINGS = {
    "page": {
        "width_mm": 210,
        "height_mm": 297,
        "margins": {"top": 2.8, "bottom": 2.2, "left": 3.0, "right": 2.0},
        "header_distance": 1.8,
        "footer_distance": 1.4,
    },
    "fonts": {
        "body_cjk": "宋体",
        "body_latin": "Times New Roman",
        "heading_cjk": "黑体",
        "heading_latin": "Times New Roman",
        "cover_title_cjk": "黑体",
        "cover_info_cjk": "宋体",
    },
    "sizes": {
        "body": 12,       # 小四
        "h1": 16,         # 三号
        "h2": 15,         # 小三
        "h3": 14,         # 四号
        "h4": 12,         # 小四
        "cover_title": 22,  # 二号
        "cover_info": 14,   # 四号
        "caption": 10.5,    # 五号
    },
    "line_spacing": 1.5,
    "first_indent_chars": 2,
}


def _load_spec(spec_path: str | Path | None) -> dict | None:
    if not spec_path:
        return None
    p = Path(spec_path)
    if not p.exists():
        return None
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return None


def _get(spec: dict | None, *keys, default=None):
    """Safely traverse nested dict."""
    if not spec:
        return default
    current = spec
    for k in keys:
        if isinstance(current, dict) and k in current:
            current = current[k]
        else:
            return default
    return current


# -------------------------------------------------------------------
# Style setup
# -------------------------------------------------------------------

def _setup_styles(doc, settings: dict):
    """Configure all document styles."""
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    fonts = settings["fonts"]
    sizes = settings["sizes"]

    # --- Normal style ---
    normal = doc.styles["Normal"]
    normal.font.size = Pt(sizes["body"])
    # Set line spacing using MULTIPLE rule (so python-docx reads it back as a ratio)
    from docx.enum.text import WD_LINE_SPACING
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = settings["line_spacing"]
    normal.paragraph_format.first_line_indent = Pt(sizes["body"] * settings["first_indent_chars"])
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    # Set CJK + Latin fonts via XML
    rPr = normal.font._element
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), fonts["body_cjk"])
    rFonts.set(qn("w:ascii"), fonts["body_latin"])
    rFonts.set(qn("w:hAnsi"), fonts["body_latin"])

    # --- Heading styles ---
    heading_configs = [
        (1, sizes["h1"], fonts["heading_cjk"]),
        (2, sizes["h2"], fonts["heading_cjk"]),
        (3, sizes["h3"], fonts["heading_cjk"]),
        (4, sizes["h4"], fonts["heading_cjk"]),
    ]

    for level, size_pt, cjk_font in heading_configs:
        try:
            style = doc.styles[f"Heading {level}"]
        except KeyError:
            continue

        style.font.size = Pt(size_pt)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.first_line_indent = None
        style.paragraph_format.alignment = None  # Left-aligned (default)

        # Set fonts
        rPr = style.font._element
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), cjk_font)
        rFonts.set(qn("w:ascii"), fonts["heading_latin"])
        rFonts.set(qn("w:hAnsi"), fonts["heading_latin"])


def _setup_page(doc, settings: dict):
    """Configure page setup."""
    from docx.shared import Cm

    section = doc.sections[0]
    margins = settings["page"]["margins"]

    section.page_width = Cm(settings["page"]["width_mm"] / 10)
    section.page_height = Cm(settings["page"]["height_mm"] / 10)
    section.top_margin = Cm(margins["top"])
    section.bottom_margin = Cm(margins["bottom"])
    section.left_margin = Cm(margins["left"])
    section.right_margin = Cm(margins["right"])

    hd = settings["page"].get("header_distance")
    fd = settings["page"].get("footer_distance")
    if hd:
        section.header_distance = Cm(hd)
    if fd:
        section.footer_distance = Cm(fd)


# -------------------------------------------------------------------
# Content helpers
# -------------------------------------------------------------------

def _add_heading(doc, text: str, level: int = 1, center: bool = False):
    """Add a heading paragraph."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_heading(text, level=level)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _add_body(doc, text: str):
    """Add a body paragraph."""
    return doc.add_paragraph(text)


def _add_centered(doc, text: str, size_pt: int = None, bold: bool = False, cjk_font: str = None):
    """Add a centered paragraph with optional formatting."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    if size_pt:
        run.font.size = Pt(size_pt)
    if bold:
        run.font.bold = True
    if cjk_font:
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), cjk_font)
    return p


def _add_superscript_ref(doc, text: str, ref_number: int):
    """Add a paragraph with a superscript citation reference."""
    from docx.shared import Pt
    p = doc.add_paragraph()
    run_text = p.add_run(text)
    run_ref = p.add_run(f"[{ref_number}]")
    run_ref.font.superscript = True
    run_ref.font.size = Pt(10)
    return p


def _add_three_line_table(doc, headers: list[str], data: list[list[str]] = None):
    """Add a three-line table (top/bottom thick, header bottom thin, no verticals)."""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    rows = 1 + (len(data) if data else 0)
    table = doc.add_table(rows=rows, cols=len(headers))

    # Set table borders: three-line style
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    borders = OxmlElement("w:tblBorders")
    for side in ("top", "bottom"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "12")  # 1.5pt
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    for side in ("left", "right", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    # Header bottom: thin
    insideH = OxmlElement("w:insideH")
    insideH.set(qn("w:val"), "single")
    insideH.set(qn("w:sz"), "6")  # 0.75pt
    insideH.set(qn("w:space"), "0")
    insideH.set(qn("w:color"), "000000")
    borders.append(insideH)
    tblPr.append(borders)

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10.5)

    # Data rows
    if data:
        for row_idx, row_data in enumerate(data):
            for col_idx, val in enumerate(row_data):
                table.rows[row_idx + 1].cells[col_idx].text = val

    return table


# -------------------------------------------------------------------
# Template generators
# -------------------------------------------------------------------

def generate_journal_template(doc, settings: dict):
    """Generate journal paper template structure."""
    from docx.shared import Pt

    sizes = settings["sizes"]
    fonts = settings["fonts"]

    # --- Title ---
    _add_centered(doc, "论文标题", size_pt=sizes["h1"], bold=True, cjk_font=fonts["heading_cjk"])
    _add_centered(doc, "作者姓名", size_pt=sizes["body"])
    _add_centered(doc, "（学校名称 学院，城市 邮编）", size_pt=10.5)

    doc.add_paragraph("")

    # --- Abstract ---
    _add_centered(doc, "摘  要", size_pt=sizes["h1"], bold=True, cjk_font=fonts["heading_cjk"])
    _add_body(doc, "在此输入中文摘要内容。摘要应概括论文的主要内容，包括研究目的、方法、结果和结论，一般不超过300字。")
    _add_body(doc, "")

    # Keywords
    p_kw = doc.add_paragraph()
    run_label = p_kw.add_run("关键词：")
    run_label.font.bold = True
    run_label.font.size = Pt(sizes["body"])
    run_kw = p_kw.add_run("关键词1；关键词2；关键词3；关键词4")
    run_kw.font.size = Pt(sizes["body"])

    doc.add_paragraph("")

    # --- English Abstract ---
    _add_centered(doc, "Abstract", size_pt=sizes["h1"], bold=True)
    _add_body(doc, "Enter the English abstract here. The abstract should summarize the main content of the paper, including the purpose, methods, results, and conclusions.")

    p_kw_en = doc.add_paragraph()
    run_label_en = p_kw_en.add_run("Keywords: ")
    run_label_en.font.bold = True
    run_label_en.font.size = Pt(sizes["body"])
    run_kw_en = p_kw_en.add_run("keyword1; keyword2; keyword3; keyword4")
    run_kw_en.font.size = Pt(sizes["body"])

    doc.add_page_break()

    # --- Body ---
    _add_heading(doc, "1 引言", level=1)
    _add_body(doc, "在此输入引言内容。引言应说明研究背景、研究目的和意义，以及论文的组织结构。")

    _add_heading(doc, "1.1 研究背景", level=2)
    _add_body(doc, "在此输入研究背景。可引用参考文献，如：研究表明该领域具有重要意义[1]。")

    _add_heading(doc, "2 研究方法", level=1)
    _add_body(doc, "在此描述研究方法。")

    _add_heading(doc, "2.1 数据来源", level=2)
    _add_body(doc, "在此描述数据来源和采集方法。")

    _add_heading(doc, "2.2 分析方法", level=2)
    _add_body(doc, "在此描述分析方法。可插入表格：")

    # Example table
    doc.add_paragraph("")
    _add_centered(doc, "表 1 示例表格", size_pt=10.5, cjk_font=fonts["body_cjk"])
    _add_three_line_table(doc,
        headers=["项目", "指标A", "指标B", "指标C"],
        data=[
            ["样本1", "0.85", "0.92", "0.78"],
            ["样本2", "0.91", "0.88", "0.83"],
            ["样本3", "0.79", "0.95", "0.87"],
        ]
    )
    doc.add_paragraph("")

    _add_heading(doc, "3 结果与分析", level=1)
    _add_body(doc, "在此展示研究结果并进行分析。")

    _add_heading(doc, "4 讨论", level=1)
    _add_body(doc, "在此讨论研究结果的意义和局限性。")

    _add_heading(doc, "5 结论", level=1)
    _add_body(doc, "在此总结论文的主要发现和贡献。")

    doc.add_page_break()

    # --- References ---
    _add_heading(doc, "参考文献", level=1)
    refs = [
        "[1] 作者：《书名》，出版地：出版社，年份，第X页。",
        "[2] Author, \"Article Title,\" Journal Name, vol. X, no. X (Year), pp. X-X.",
        "[3] 作者：《文章题名》，《期刊名》年份第X期。",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Pt(24)  # hanging indent
        # hanging indent via first-line negative
        pPr = p._p.get_or_add_pPr()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        ind.set(qn("w:left"), "480")  # 2 chars
        ind.set(qn("w:hanging"), "480")


def generate_thesis_template(doc, settings: dict):
    """Generate thesis (graduation paper) template structure."""
    from docx.shared import Pt

    sizes = settings["sizes"]
    fonts = settings["fonts"]

    # --- Cover page ---
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")

    _add_centered(doc, "XX 大学", size_pt=26, bold=True, cjk_font=fonts["cover_title_cjk"])
    doc.add_paragraph("")
    _add_centered(doc, "本科毕业论文（设计）", size_pt=22, bold=True, cjk_font=fonts["cover_title_cjk"])

    doc.add_paragraph("")
    doc.add_paragraph("")

    _add_centered(doc, "论文题目", size_pt=sizes["h1"], bold=True, cjk_font=fonts["heading_cjk"])
    doc.add_paragraph("")

    # Student info
    info_items = [
        ("学    院", "XX学院"),
        ("专    业", "XX专业"),
        ("学    号", "20XXXXXXXX"),
        ("姓    名", "XXX"),
        ("指导教师", "XXX 教授"),
        ("完成日期", "2026年XX月"),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = 1  # CENTER
        run_l = p.add_run(f"{label}：")
        run_l.font.size = Pt(sizes["cover_info"])
        run_l.font.bold = True
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        rPr = run_l._r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), fonts["cover_info_cjk"])

        run_v = p.add_run(f"  {value}  ")
        run_v.font.size = Pt(sizes["cover_info"])
        # Underline
        from docx.oxml import OxmlElement as OE
        rPr2 = run_v._r.get_or_add_rPr()
        u = OE("w:u")
        u.set(qn("w:val"), "single")
        rPr2.append(u)

    doc.add_page_break()

    # --- Chinese Abstract ---
    _add_centered(doc, "摘  要", size_pt=sizes["h1"], bold=True, cjk_font=fonts["heading_cjk"])
    _add_body(doc, "在此输入中文摘要内容。摘要应概括论文的研究目的、方法、结果和主要结论。学位论文摘要一般为500-1000字。")
    _add_body(doc, "")

    p_kw = doc.add_paragraph()
    run_label = p_kw.add_run("关键词：")
    run_label.font.bold = True
    run_label.font.size = Pt(sizes["body"])
    p_kw.add_run("关键词1；关键词2；关键词3；关键词4；关键词5").font.size = Pt(sizes["body"])

    doc.add_page_break()

    # --- English Abstract ---
    _add_centered(doc, "Abstract", size_pt=sizes["h1"], bold=True)
    _add_body(doc, "Enter the English abstract here. The abstract should summarize the research purpose, methods, results, and main conclusions of the thesis.")

    p_kw_en = doc.add_paragraph()
    run_label = p_kw_en.add_run("Keywords: ")
    run_label.font.bold = True
    p_kw_en.add_run("keyword1; keyword2; keyword3; keyword4; keyword5")

    doc.add_page_break()

    # --- Table of Contents placeholder ---
    _add_centered(doc, "目  录", size_pt=sizes["h1"], bold=True, cjk_font=fonts["heading_cjk"])
    _add_body(doc, "（此处由 Word 自动生成目录：引用 → 目录 → 自动目录）")
    _add_body(doc, "")

    doc.add_page_break()

    # --- Body chapters ---
    _add_heading(doc, "第一章 绪论", level=1, center=False)
    _add_heading(doc, "1.1 研究背景", level=2)
    _add_body(doc, "在此介绍研究背景和意义。")
    _add_heading(doc, "1.2 国内外研究现状", level=2)
    _add_body(doc, "在此综述相关领域的研究现状。")
    _add_heading(doc, "1.3 研究目的与意义", level=2)
    _add_body(doc, "在此说明本文的研究目的和理论/实践意义。")
    _add_heading(doc, "1.4 论文结构安排", level=2)
    _add_body(doc, "在此介绍论文各章节的内容安排。")

    doc.add_page_break()

    _add_heading(doc, "第二章 相关理论与方法", level=1, center=False)
    _add_heading(doc, "2.1 理论基础", level=2)
    _add_body(doc, "在此介绍论文涉及的核心理论。")
    _add_heading(doc, "2.2 研究方法", level=2)
    _add_body(doc, "在此介绍论文采用的研究方法。")

    doc.add_page_break()

    _add_heading(doc, "第三章 研究内容", level=1, center=False)
    _add_heading(doc, "3.1 研究设计", level=2)
    _add_body(doc, "在此描述研究设计。")
    _add_heading(doc, "3.2 数据分析", level=2)
    _add_body(doc, "在此展示数据分析结果。")

    doc.add_page_break()

    _add_heading(doc, "第四章 结果与讨论", level=1, center=False)
    _add_heading(doc, "4.1 研究结果", level=2)
    _add_body(doc, "在此展示研究结果。")
    _add_heading(doc, "4.2 讨论", level=2)
    _add_body(doc, "在此讨论研究结果的意义。")

    doc.add_page_break()

    _add_heading(doc, "第五章 结论与展望", level=1, center=False)
    _add_heading(doc, "5.1 主要结论", level=2)
    _add_body(doc, "在此总结论文的主要发现和结论。")
    _add_heading(doc, "5.2 不足与展望", level=2)
    _add_body(doc, "在此说明研究的局限性和未来研究方向。")

    doc.add_page_break()

    # --- References ---
    _add_centered(doc, "参考文献", size_pt=sizes["h1"], bold=True, cjk_font=fonts["heading_cjk"])
    refs = [
        "[1] 作者：《书名》，出版地：出版社，年份，第X页。",
        "[2] Author, *Title*, Place: Publisher, Year, p. X.",
        "[3] 作者：《文章题名》，《期刊名》年份第X期。",
        "[4] Author, \"Article,\" *Journal*, vol. X, no. X (Year), pp. X-X.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.first_line_indent = None
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = p._p.get_or_add_pPr()
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        ind.set(qn("w:left"), "480")
        ind.set(qn("w:hanging"), "480")

    doc.add_page_break()

    # --- Acknowledgments ---
    _add_centered(doc, "致  谢", size_pt=sizes["h1"], bold=True, cjk_font=fonts["heading_cjk"])
    _add_body(doc, "在此撰写致谢内容。感谢导师、同学、家人等在论文写作过程中给予的帮助和支持。")

    doc.add_page_break()

    # --- Appendices ---
    _add_centered(doc, "附  录", size_pt=sizes["h1"], bold=True, cjk_font=fonts["heading_cjk"])
    _add_heading(doc, "附录A 数据表格", level=2)
    _add_body(doc, "在此放置补充数据表格。")


# -------------------------------------------------------------------
# Main
# -------------------------------------------------------------------

def generate_template(
    mode: str = "journal",
    spec_path: str | Path | None = None,
    output_path: str | Path | None = None,
    spec: dict | None = None,
) -> str:
    """Generate a formatted .docx template.

    Args:
        mode: "journal" or "thesis"
        spec_path: Path to spec JSON from parse_spec.py
        output_path: Output .docx path
        spec: Pre-loaded spec dict

    Returns path to generated file.
    """
    from docx import Document

    # Load spec
    if spec is None and spec_path:
        spec = _load_spec(spec_path)

    # Build settings from spec or defaults
    settings = dict(DEFAULT_SETTINGS)
    if spec:
        margins = _get(spec, "page_setup", "margins")
        if margins:
            for side in ("top", "bottom", "left", "right"):
                if side in margins:
                    settings["page"]["margins"][side] = margins[side] / 10  # mm -> cm

        for level in range(1, 5):
            h_size = _get(spec, "styles", f"Heading {level}", "size")
            if h_size:
                settings["sizes"][f"h{level}"] = round(h_size)

        n_size = _get(spec, "styles", "Normal", "size")
        if n_size:
            settings["sizes"]["body"] = round(n_size)

        ls = _get(spec, "paragraphs_summary", "line_spacing")
        if ls:
            settings["line_spacing"] = round(ls, 2)

        cjk = _get(spec, "styles", "Normal", "font_cjk")
        if cjk:
            settings["fonts"]["body_cjk"] = cjk
        latin = _get(spec, "styles", "Normal", "font_latin")
        if latin:
            settings["fonts"]["body_latin"] = latin

    # Create document
    doc = Document()

    # Setup styles and page
    _setup_styles(doc, settings)
    _setup_page(doc, settings)

    # Generate content
    if mode == "thesis":
        generate_thesis_template(doc, settings)
    else:
        generate_journal_template(doc, settings)

    # Save
    if output_path is None:
        output_path = f"论文模板_{mode}.docx"

    output_path = Path(output_path)
    doc.save(str(output_path))
    return str(output_path)


def main():
    from scripts.lib import setup_utf8_output
    setup_utf8_output()

    parser = argparse.ArgumentParser(description="Generate Chinese academic paper template")
    parser.add_argument(
        "--mode", "-m", choices=["journal", "thesis"], default="journal",
        help="Template mode: journal (default) or thesis",
    )
    parser.add_argument("--spec", "-s", help="Path to spec JSON for custom formatting")
    parser.add_argument("--output", "-o", help="Output .docx path")
    args = parser.parse_args()

    output = generate_template(
        mode=args.mode,
        spec_path=args.spec,
        output_path=args.output,
    )
    print(f"Template generated: {output}")


if __name__ == "__main__":
    main()
