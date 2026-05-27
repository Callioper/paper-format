#!/usr/bin/env python3
"""generate_report.py - Generate unified comparison report (.docx).

Combines:
  1. paper-check format comparison report (original → repaired)
  2. cite-formatter reference format report section
"""

from __future__ import annotations
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# -------------------------------------------------------------------
# Helpers
# -------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_heading_row(table, text: str, level: int = 1):
    """Add a section heading row to a table."""
    row = table.add_row()
    cell = row.cells[0]
    cell.text = text
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(11 if level > 1 else 12)
    _set_cell_bg(cell, "D9E1F2")
    return row


# -------------------------------------------------------------------
# Main report generation
# -------------------------------------------------------------------

def generate_unified_report(
    original_path: str | Path,
    repaired_path: str | Path,
    format_issues: dict,
    citation_results: list[dict],
    output_path: str | Path,
    footnote_results: list[dict] | None = None,
    repair_records: list[dict] | None = None,
    verification_results: dict | None = None,
    quote_results: list[dict] | None = None,
    mode: str = "thesis",
) -> None:
    """
    Generate a unified A4 landscape comparison report.

    Sections:
      1. Format issues (margins, spacing, fonts -- with location + before/after)
      2. Citation format report
      3. Footnote citation report
      4. Citation verification report (if verification data provided)
      5. Repair statistics summary
      6. Quote verification report (if quote_results provided)
    """
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.orientation = 1  # landscape

    # Title
    title = doc.add_heading("论文格式检测与修复报告", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta = doc.add_paragraph()
    meta.add_run(f"原始文件：{Path(original_path).name}").bold = False
    meta.add_run(f"\n修复后文件：{Path(repaired_path).name}")
    meta.add_run(f"\n文档模式：{'学位论文' if mode == 'thesis' else '期刊论文'}")
    meta.add_run(f"\n格式规范：CNU《外国文学评论》2024修订版 + GB/T 7713.1-2006")

    doc.add_paragraph("")

    # -----------------------------------------------------------------
    # Section 1: Format issues (with location + before/after)
    # -----------------------------------------------------------------
    doc.add_heading("一、格式问题", level=2)
    issues = format_issues.get("issues", {})
    if any(issues.values()):
        for category, items in issues.items():
            if not items:
                continue
            doc.add_heading(f"  {category}", level=3)
            table = doc.add_table(rows=1, cols=6)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            for cell, text in zip(hdr, ["序号", "位置", "检查项", "规范要求", "当前值", "修复建议"]):
                cell.text = text
                cell.paragraphs[0].runs[0].bold = True
                _set_cell_bg(cell, "D9E1F2")
            for i, item in enumerate(items, 1):
                row = table.add_row().cells
                row[0].text = str(i)
                row[1].text = item.get("location", item.get("sample_text", ""))
                row[2].text = item.get("item", "")
                row[3].text = item.get("expected", "")
                row[4].text = item.get("actual", "")
                row[5].text = item.get("suggestion", "")
    else:
        doc.add_paragraph("未发现格式问题")

    doc.add_paragraph("")

    # -----------------------------------------------------------------
    # Section 1b: Repair records (before/after comparison)
    # -----------------------------------------------------------------
    if repair_records:
        doc.add_heading("  修复记录（修复前 → 修复后）", level=3)
        rtable = doc.add_table(rows=1, cols=5)
        rtable.style = "Table Grid"
        rhdr = rtable.rows[0].cells
        for cell, text in zip(rhdr, ["类别", "项目", "位置", "修复前", "修复后"]):
            cell.text = text
            cell.paragraphs[0].runs[0].bold = True
            _set_cell_bg(cell, "E2EFDA")
        for rec in repair_records:
            if rec.get("error"):
                continue
            row = rtable.add_row().cells
            row[0].text = rec.get("category", "")
            row[1].text = rec.get("item", "")
            row[2].text = rec.get("location", "")
            before = rec.get("before", "")
            after = rec.get("after", "")
            row[3].text = str(before) if before else ""
            row[4].text = str(after) if after else ""
        doc.add_paragraph("")

    # -----------------------------------------------------------------
    # Section 2: Citation format report (from cite-formatter)
    # -----------------------------------------------------------------
    doc.add_heading("二、参考文献格式报告", level=2)

    if not citation_results:
        doc.add_paragraph("✅ 未检测到参考文献章节")
    else:
        total = len(citation_results)
        p1_count = sum(1 for r in citation_results if any("P1" in w for w in r.get("warnings", [])))
        p2_count = sum(1 for r in citation_results if any("P2" in w for w in r.get("warnings", [])))
        p3_count = sum(1 for r in citation_results if any("P3" in w for w in r.get("warnings", [])))

        doc.add_paragraph(f"共检测 {total} 条引用")
        doc.add_paragraph(
            f"问题分布：P1(必须修复) {p1_count} 条，P2(应当修复) {p2_count} 条，P3(建议修复) {p3_count} 条"
        )

        doc.add_paragraph("")

        for idx, r in enumerate(citation_results, 1):
            # Original
            p_orig = doc.add_paragraph()
            p_orig.add_run(f"[{idx}] 原文：").bold = True
            p_orig.add_run(r.get("original", ""))

            # Formatted
            formatted = r.get("formatted", "")
            if formatted:
                p_fmt = doc.add_paragraph()
                p_fmt.add_run("     规范化：").bold = True
                p_fmt.add_run(formatted)

            # Warnings
            warnings = r.get("warnings", [])
            for w in warnings:
                p_w = doc.add_paragraph()
                p_w.add_run(f"     {w}")
                if "P1" in w:
                    p_w.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                elif "P2" in w:
                    p_w.runs[0].font.color.rgb = RGBColor(0xFF, 0x66, 0x00)

            # BibTeX status
            bib = r.get("bib_status", r.get("zotero_status", ""))
            if bib and bib != "not_checked":
                doc.add_paragraph(f"     BibTeX状态：{bib}")

            # Manual review flag
            if r.get("needs_manual_review"):
                p_manual = doc.add_paragraph()
                p_manual.add_run("     ⚠️ 需要人工核实").bold = True
                p_manual.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

            doc.add_paragraph("")

    # -----------------------------------------------------------------
    # Section 3: Footnote citation report
    # -----------------------------------------------------------------
    doc.add_heading("三、脚注格式报告", level=2)

    if not footnote_results:
        doc.add_paragraph("✅ 未检测到脚注引用问题")
    else:
        fn_with_warns = [r for r in footnote_results if r.get("warnings")]
        fn_ok = len(footnote_results) - len(fn_with_warns)
        p1_fn = sum(1 for r in footnote_results if any("P1" in w for w in r.get("warnings", [])))
        p2_fn = sum(1 for r in footnote_results if any("P2" in w for w in r.get("warnings", [])))

        doc.add_paragraph(f"共检测 {len(footnote_results)} 条脚注引用，{fn_ok} 条通过，{len(fn_with_warns)} 条有问题")
        doc.add_paragraph(f"问题分布：P1(必须修复) {p1_fn} 条，P2(应当修复) {p2_fn} 条")
        doc.add_paragraph("")

        for r in footnote_results:
            fn_id = r.get("footnote_id", "?")
            raw = r.get("raw_footnote_text", r.get("original", ""))
            formatted = r.get("formatted", "")
            warnings = r.get("warnings", [])

            p_orig = doc.add_paragraph()
            p_orig.add_run(f"脚注[{fn_id}] 原文：").bold = True
            p_orig.add_run(raw[:120])

            if formatted and warnings:
                p_fmt = doc.add_paragraph()
                p_fmt.add_run("     规范化：").bold = True
                p_fmt.add_run(formatted)

            for w in warnings:
                p_w = doc.add_paragraph()
                p_w.add_run(f"     {w}")
                if "P1" in w:
                    p_w.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                elif "P2" in w:
                    p_w.runs[0].font.color.rgb = RGBColor(0xFF, 0x66, 0x00)

            if r.get("needs_manual_review"):
                pm = doc.add_paragraph()
                pm.add_run("     ⚠️ 需要人工核实").bold = True
                pm.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

            doc.add_paragraph("")

    # -----------------------------------------------------------------
    # Section 4: Citation verification report
    # -----------------------------------------------------------------
    if verification_results:
        doc.add_heading("四、引文完整性校验报告", level=2)

        # Local verification
        local = verification_results.get("local", {})
        if local:
            doc.add_heading("  本地数据库验证", level=3)
            doc.add_paragraph(
                f"数据库来源：{local.get('db_source', 'N/A')}，"
                f"条目数：{local.get('db_entries', 0)}，"
                f"匹配：{local.get('matched', 0)}/{local.get('citations_total', 0)}"
            )
            unmatched = [r for r in local.get("results", []) if not r.get("local_match")]
            if unmatched:
                doc.add_paragraph(f"未匹配条目（{len(unmatched)} 条）：")
                for r in unmatched[:20]:
                    p = doc.add_paragraph()
                    p.add_run(f"  [{r.get('citation_index', '?')}] ").bold = True
                    p.add_run(r.get("raw_text", "")[:80])
                    for issue in r.get("issues", []):
                        doc.add_paragraph(f"      {issue}")

        # External verification
        ext = verification_results.get("external", {})
        if ext:
            doc.add_heading("  外部数据源验证", level=3)
            sources_used = ", ".join(ext.get("sources_used", []))
            doc.add_paragraph(
                f"数据源：{sources_used}，"
                f"匹配：{ext.get('matched', 0)}/{ext.get('citations_total', 0)}"
            )
            unmatched_ext = [r for r in ext.get("results", []) if not r.get("external_match")]
            if unmatched_ext:
                doc.add_paragraph(f"未匹配条目（{len(unmatched_ext)} 条）：")
                for r in unmatched_ext[:20]:
                    p = doc.add_paragraph()
                    p.add_run(f"  [{r.get('citation_index', '?')}] ").bold = True
                    p.add_run(r.get("raw_text", "")[:80])

        # Cross-check
        cross = verification_results.get("cross_check", {})
        if cross:
            doc.add_heading("  交叉校验", level=3)
            orphans = cross.get("orphan_citations", [])
            unused = cross.get("unused_references", [])
            dupes = cross.get("duplicates", [])
            incomplete = cross.get("incomplete", [])

            total_issues = len(orphans) + len(unused) + len(dupes) + len(incomplete)
            if total_issues == 0:
                doc.add_paragraph("交叉校验通过，无异常")
            else:
                doc.add_paragraph(f"发现 {total_issues} 个问题")
                for item in orphans[:10]:
                    p = doc.add_paragraph()
                    p.add_run("  [孤立引用] ").font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                    p.add_run(item.get("issue", ""))
                for item in unused[:10]:
                    p = doc.add_paragraph()
                    p.add_run("  [未使用文献] ").font.color.rgb = RGBColor(0xFF, 0x66, 0x00)
                    p.add_run(item.get("issue", ""))
                for item in dupes[:10]:
                    p = doc.add_paragraph()
                    p.add_run("  [重复引用] ").font.color.rgb = RGBColor(0xFF, 0x66, 0x00)
                    p.add_run(item.get("issue", ""))
                for item in incomplete[:10]:
                    p = doc.add_paragraph()
                    p.add_run("  [不完整] ").font.color.rgb = RGBColor(0xFF, 0x66, 0x00)
                    p.add_run(item.get("issue", ""))

    # -----------------------------------------------------------------
    # Section 5: Statistics summary
    # -----------------------------------------------------------------
    doc.add_heading("五、统计摘要", level=2)
    stats_table = doc.add_table(rows=1, cols=3)
    stats_table.style = "Table Grid"
    shdr = stats_table.rows[0].cells
    for cell, text in zip(shdr, ["类别", "问题数", "说明"]):
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True
        _set_cell_bg(cell, "D9E1F2")

    category_names = {
        "page_setup": "页面设置",
        "styles": "样式",
        "paragraphs": "段落",
        "tables": "表格",
        "cover": "封面",
        "abstract": "摘要",
        "toc": "目录",
        "headers_footers": "页眉页脚",
        "acknowledgments": "致谢",
        "appendices": "附录",
        "citations": "参考文献",
        "footnotes": "脚注",
    }
    for cat, items in issues.items():
        if not items:
            continue
        row = stats_table.add_row().cells
        row[0].text = category_names.get(cat, cat)
        row[1].text = str(len(items))
        row[2].text = ", ".join(set(item.get("severity", "") for item in items))

    # -----------------------------------------------------------------
    # Section 6: Quote verification report
    # -----------------------------------------------------------------
    if quote_results:
        doc.add_heading("六、引文逐字核对报告", level=2)

        total_q = len(quote_results)
        status_counts = {}
        for r in quote_results:
            s = r.get("status", "")
            status_counts[s] = status_counts.get(s, 0) + 1

        summary_parts = []
        for s in ("一致", "基本一致", "疑似不一致", "未定位到出处"):
            if s in status_counts:
                summary_parts.append(f"{s} {status_counts[s]} 条")

        doc.add_paragraph(f"共核对 {total_q} 条直接引文：{', '.join(summary_parts)}")
        doc.add_paragraph("")

        # Verification table
        qtable = doc.add_table(rows=1, cols=6)
        qtable.style = "Table Grid"
        qhdr = qtable.rows[0].cells
        for cell, text in zip(qhdr, ["序号", "论文中的引文", "出处文件", "页码", "核对结果", "备注"]):
            cell.text = text
            cell.paragraphs[0].runs[0].bold = True
            _set_cell_bg(cell, "D9E1F2")

        for r in quote_results:
            row = qtable.add_row().cells
            row[0].text = str(r.get("quote_index", ""))
            row[1].text = r.get("quote_text", "")[:60]
            row[2].text = r.get("source_file", "")
            row[3].text = str(r.get("matched_page", "")) if r.get("matched_page") else ""

            status = r.get("status", "")
            row[4].text = status
            # Color code status
            if status == "一致":
                row[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            elif status == "基本一致":
                row[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0x99, 0x00)
            elif status == "疑似不一致":
                row[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

            row[5].text = r.get("note", "")

        # Embed screenshots if available
        screenshots_exist = any(r.get("screenshot_path") for r in quote_results)
        if screenshots_exist:
            doc.add_paragraph("")
            doc.add_heading("  出处截图证据", level=3)
            for r in quote_results:
                ss_path = r.get("screenshot_path", "")
                if ss_path and Path(ss_path).exists():
                    p = doc.add_paragraph()
                    p.add_run(f"引文[{r.get('quote_index', '?')}] 出处第{r.get('matched_page', '?')}页：").bold = True
                    try:
                        doc.add_picture(ss_path, width=Cm(16))
                    except Exception:
                        doc.add_paragraph(f"  [截图: {ss_path}]")

    doc.save(output_path)
    print(f"[generate_report] Report saved to {output_path}")


# -------------------------------------------------------------------
# Standalone citation-only report (cite-formatter style)
# -------------------------------------------------------------------

def generate_citation_report_only(
    citation_results: list[dict],
    output_path: str | Path,
) -> None:
    """Generate a standalone citation format report."""
    doc = Document()
    doc.add_heading("参考文献格式报告", level=1)

    total = len(citation_results)
    p1 = sum(1 for r in citation_results if any("P1" in w for w in r.get("warnings", [])))
    p2 = sum(1 for r in citation_results if any("P2" in w for w in r.get("warnings", [])))
    p3 = sum(1 for r in citation_results if any("P3" in w for w in r.get("warnings", [])))

    doc.add_paragraph(f"共检测 {total} 条引用")
    doc.add_paragraph(f"P1(必须修复) {p1} 条，P2(应当修复) {p2} 条，P3(建议修复) {p3} 条")
    doc.add_paragraph("")

    for idx, r in enumerate(citation_results, 1):
        doc.add_paragraph(f"[{idx}] {r.get('original', '')[:80]}")
        if r.get("formatted"):
            doc.add_paragraph(f"  → {r['formatted']}")
        for w in r.get("warnings", []):
            doc.add_paragraph(f"  {w}")

    doc.save(output_path)


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser()
    parser.add_argument("original", help="Original thesis .docx")
    parser.add_argument("repaired", help="Repaired thesis .docx")
    parser.add_argument("output", help="Output report .docx path")
    parser.add_argument("--format-issues", help="JSON file with format issues")
    parser.add_argument("--citations", help="JSON file with citation results")
    args = parser.parse_args()

    format_issues = {}
    if args.format_issues:
        import json
        format_issues = json.loads(Path(args.format_issues).read_text(encoding="utf-8"))

    citation_results = []
    if args.citations:
        import json
        citation_results = json.loads(Path(args.citations).read_text(encoding="utf-8"))

    generate_unified_report(
        args.original,
        args.repaired,
        format_issues,
        citation_results,
        args.output,
    )