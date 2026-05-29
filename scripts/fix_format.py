#!/usr/bin/env python3
"""fix_format.py - Automatic format repair for thesis documents.

Supports two document modes:
  - thesis:  cover, abstract, TOC, body, references, acknowledgments, appendices
  - journal: title page, abstract, keywords, body, references, footnotes

Supports spec-driven repair via parse_spec.py output JSON.
Records before/after values for each repair action.
"""

from __future__ import annotations
import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

_THIS = Path(__file__).resolve()
_REPO_ROOT = _THIS.parents[1]
if str(_REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(_REPO_ROOT))

from scripts.extract_references import detect_references_section
from scripts.citation_repair import process_citations, repair_citations_in_docx


# -------------------------------------------------------------------
# Document modes (mirrors check_format.py)
# -------------------------------------------------------------------

DOC_MODES = {
    "thesis": ["cover", "abstract", "toc", "body", "references", "acknowledgments", "appendices"],
    "journal": ["title_page", "abstract", "keywords", "body", "references", "footnotes"],
}


# -------------------------------------------------------------------
# Default repair targets
# -------------------------------------------------------------------

DEFAULT_HEADING_SIZES = {1: 16, 2: 15, 3: 14, 4: 12}
DEFAULT_MARGINS = {"top": 2.8, "bottom": 2.2, "left": 3.0, "right": 2.0}
DEFAULT_BODY_SIZE = 12
DEFAULT_LINE_SPACING = 1.5
DEFAULT_FIRST_LINE_INDENT_PT = 24  # 2 chars at 12pt


def _load_spec(spec_path: str | Path | None) -> dict | None:
    """Load a spec JSON file produced by parse_spec.py."""
    if not spec_path:
        return None
    p = Path(spec_path)
    if not p.exists():
        return None
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return None


def _get_spec_value(spec: dict | None, *keys, default=None):
    """Safely traverse nested spec dict."""
    if not spec:
        return default
    current = spec
    for k in keys:
        if isinstance(current, dict) and k in current:
            current = current[k]
        else:
            return default
    return current


# -------------------------------------------------------------------
# Font / paragraph helpers
# -------------------------------------------------------------------

def _pt(length) -> float | None:
    """Convert a python-docx Length to pt, or return None."""
    try:
        return length.pt if length is not None else None
    except Exception:
        return None


def _apply_font_to_run(run, chinese="宋体", english="Times New Roman", size_pt=12):
    """Set CJK + Latin fonts and size on a single run via XML."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt

    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), chinese)
    rFonts.set(qn("w:ascii"), english)
    rFonts.set(qn("w:hAnsi"), english)
    run.font.size = Pt(size_pt)


# CJK faces that carry meaning and must survive body normalization. In Chinese
# academic typesetting, block quotes / epigraphs are conventionally set in 楷体
# (sometimes 仿宋); flattening them to 宋体 destroys an intentional distinction the
# author made. So for these runs we keep the face and only correct the size.
_PRESERVE_CJK = ("楷", "仿宋")


def _apply_body_font(para, chinese="宋体", english="Times New Roman", size=12):
    from docx.oxml.ns import qn
    from docx.shared import Pt
    for run in para.runs:
        rPr = run._r.find(qn("w:rPr"))
        cur = None
        if rPr is not None:
            rf = rPr.find(qn("w:rFonts"))
            if rf is not None:
                cur = rf.get(qn("w:eastAsia"))
        if cur and any(k in cur for k in _PRESERVE_CJK):
            run.font.size = Pt(size)          # keep 楷体/仿宋 face, normalize size only
        else:
            _apply_font_to_run(run, chinese=chinese, english=english, size_pt=size)


def _get_run_font_size(para) -> float | None:
    """Get the font size of the first run in a paragraph."""
    for run in para.runs:
        if run.font.size:
            return _pt(run.font.size)
    return None


# -------------------------------------------------------------------
# Heading fix
# -------------------------------------------------------------------

def _fix_heading(para, heading_sizes: dict | None = None) -> bool:
    """Apply correct heading font size. Returns True if fixed."""
    sizes = heading_sizes or DEFAULT_HEADING_SIZES
    sname = para.style.name if para.style else ""
    m = re.search(r"Heading\s*(\d)", sname, re.IGNORECASE)
    if not m:
        return False
    level = int(m.group(1))
    size_pt = sizes.get(level)
    if size_pt is None:
        return False
    old_size = _get_run_font_size(para)
    if old_size is not None and abs(old_size - size_pt) < 0.5:
        return False  # Already correct
    for run in para.runs:
        _apply_font_to_run(run, chinese="黑体", english="Times New Roman", size_pt=size_pt)
    return True


# -------------------------------------------------------------------
# Section-specific fixes
# -------------------------------------------------------------------

SECTION_KEYWORDS = {
    "abstract_zh": ["摘  要", "摘 要", "摘    要", "摘 要"],
    "abstract_en": ["Abstract", "ABSTRACT"],
    "toc": ["目  录", "目 录", "目    录", "Contents", "CONTENTS"],
    "acknowledgments": ["致  谢", "致 谢", "致    谢"],
    "appendices": ["附  录", "附 录", "附    录", "Appendix", "APPENDIX"],
}


def _is_section_title(para, keyword_group: str) -> bool:
    """Check if paragraph is a section title matching keyword group."""
    text = para.text.strip()
    if not text:
        return False
    keywords = SECTION_KEYWORDS.get(keyword_group, [])
    return any(kw in text for kw in keywords)


def _fix_section_title(para, old_size_pt: float | None = None, keyword_group: str = "") -> dict:
    """Fix a section title paragraph (center, font size, font family).
    Returns repair record with before/after."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    record: dict[str, Any] = {}
    target_size = 16  # 三号
    text = para.text.strip()

    # Run-in abstract guard: many course/journal papers write the abstract as a
    # single paragraph "摘 要：<整段正文>". The label looks like a section title, but
    # the bulk of the paragraph is body text. Applying the 16pt heading size to the
    # whole paragraph (the naive `for run in para.runs` below) silently blows the
    # entire abstract up to 16pt 黑体 — a real defect we hit in practice. So when we
    # detect a label + colon followed by substantial inline content, style only the
    # label and normalize the body to 宋体 at body size instead.
    colon_idx = next((i for i, ch in enumerate(text) if ch in "：:"), -1)
    if keyword_group.startswith("abstract") and colon_idx != -1 and len(text) - colon_idx > 15:
        passed = 0
        seen_colon = False
        for run in para.runs:
            start = passed
            passed += len(run.text)
            if not seen_colon:
                _apply_font_to_run(run, chinese="黑体", english="Times New Roman", size_pt=DEFAULT_BODY_SIZE)
                run.font.bold = True
                if start <= colon_idx < passed:
                    seen_colon = True
            else:
                _apply_font_to_run(run, chinese="宋体", english="Times New Roman", size_pt=DEFAULT_BODY_SIZE)
                run.font.bold = False
        record["runin_before"] = f"{(old_size_pt or 0):.0f}pt 整段同字号"
        record["runin_after"] = f"标签黑体{DEFAULT_BODY_SIZE}pt + 正文宋体{DEFAULT_BODY_SIZE}pt"
        return record

    # Fix alignment
    old_align = para.paragraph_format.alignment
    if old_align is not None and old_align != WD_ALIGN_PARAGRAPH.CENTER:
        record["alignment_before"] = str(old_align)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        record["alignment_after"] = "CENTER"

    # Fix font size
    actual_old = old_size_pt or _get_run_font_size(para)
    if actual_old is not None and abs(actual_old - target_size) > 0.5:
        record["size_before"] = f"{actual_old:.0f}pt"
        for run in para.runs:
            _apply_font_to_run(run, chinese="黑体", english="Times New Roman", size_pt=target_size)
        record["size_after"] = f"{target_size}pt"
    elif actual_old is None:
        for run in para.runs:
            _apply_font_to_run(run, chinese="黑体", english="Times New Roman", size_pt=target_size)
        record["size_after"] = f"{target_size}pt"

    return record


# -------------------------------------------------------------------
# Main fix function
# -------------------------------------------------------------------

def fix_format(
    thesis_path: str | Path,
    spec_path: str | Path | None = None,
    output_path: str | Path | None = None,
    repair_citations: bool = True,
    mode: str = "journal",
    spec: dict | None = None,
) -> dict:
    """
    Run automatic format repair on a thesis .docx.

    Args:
        thesis_path: Path to thesis .docx
        spec_path: Path to spec JSON from parse_spec.py
        output_path: Output path for repaired file
        repair_citations: Whether to repair citation formatting
        mode: Document mode ("thesis" or "journal")
        spec: Pre-loaded spec dict (alternative to spec_path)

    Returns dict with: output, repair_records, error (if any)
    """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn
    from scripts.lib.doc_helpers import create_working_copy

    thesis_path = Path(thesis_path)
    if not thesis_path.exists():
        return {"error": f"文件不存在: {thesis_path}", "output": None, "repair_records": []}
    if thesis_path.suffix.lower() != ".docx":
        return {"error": f"仅支持 .docx 格式: {thesis_path.suffix}", "output": None, "repair_records": []}

    # Always work on a copy to preserve the original
    try:
        working_path = create_working_copy(thesis_path)
        print(f"[fix_format] Working on copy: {working_path.name}")
    except Exception as e:
        return {"error": f"无法创建副本: {e}", "output": None, "repair_records": []}

    try:
        doc = Document(working_path)
    except Exception as e:
        return {"error": f"无法打开文档: {e}", "output": None, "repair_records": []}

    # Load spec
    if spec is None and spec_path:
        spec = _load_spec(spec_path)

    # Derive repair targets from spec
    heading_sizes = dict(DEFAULT_HEADING_SIZES)
    margins_target = dict(DEFAULT_MARGINS)
    body_size = DEFAULT_BODY_SIZE
    line_spacing = DEFAULT_LINE_SPACING
    first_indent_pt = DEFAULT_FIRST_LINE_INDENT_PT

    if spec:
        for level in range(1, 5):
            h = _get_spec_value(spec, "styles", f"Heading {level}", "size")
            if h:
                heading_sizes[level] = round(h)
        spec_margins = _get_spec_value(spec, "page_setup", "margins")
        if spec_margins:
            for side in ("top", "bottom", "left", "right"):
                if side in spec_margins:
                    margins_target[side] = spec_margins[side] / 10  # mm → cm
        n_size = _get_spec_value(spec, "styles", "Normal", "size")
        if n_size:
            body_size = round(n_size)
        ls = _get_spec_value(spec, "paragraphs_summary", "line_spacing")
        if ls:
            line_spacing = round(ls, 2)
        indent = _get_spec_value(spec, "paragraphs_summary", "first_line_indent_chars")
        if indent:
            first_indent_pt = round(indent * body_size)

    sections = DOC_MODES.get(mode, DOC_MODES["thesis"])
    out_path = Path(output_path) if output_path else working_path.with_stem(working_path.stem + "_repaired")
    records: list[dict] = []

    # ---------------------------------------------------------------
    # 1. Page margins
    # ---------------------------------------------------------------
    try:
        section = doc.sections[0]
        old_margins = {}
        for side, attr in [("top", "top_margin"), ("bottom", "bottom_margin"),
                           ("left", "left_margin"), ("right", "right_margin")]:
            val = getattr(section, attr, None)
            if val is not None:
                old_margins[side] = round(val / 914400 * 25.4, 1)

        section.top_margin = Cm(margins_target["top"])
        section.bottom_margin = Cm(margins_target["bottom"])
        section.left_margin = Cm(margins_target["left"])
        section.right_margin = Cm(margins_target["right"])
        records.append({
            "category": "page_setup",
            "item": "margins",
            "location": "全文",
            "before": old_margins if old_margins else "未设置",
            "after": margins_target,
        })
    except (IndexError, AttributeError, TypeError) as e:
        records.append({"category": "page_setup", "item": "margins", "error": f"无法修改边距: {e}"})

    # ---------------------------------------------------------------
    # 2. Body paragraphs (font, line spacing, indent)
    # ---------------------------------------------------------------
    fixed_heading_count = 0
    fixed_body_count = 0
    for para in doc.paragraphs:
        try:
            name = para.style.name if para.style else ""
            if name.startswith("Heading"):
                if _fix_heading(para, heading_sizes):
                    fixed_heading_count += 1
                continue

            old_size = _get_run_font_size(para)
            old_spacing = para.paragraph_format.line_spacing
            old_indent = para.paragraph_format.first_line_indent

            _apply_body_font(para, size=body_size)
            para.paragraph_format.line_spacing = line_spacing
            try:
                para.paragraph_format.first_line_indent = Pt(first_indent_pt)
            except Exception:
                pass
            fixed_body_count += 1
        except Exception:
            pass

    records.append({
        "category": "paragraphs",
        "item": "font_indent_spacing",
        "location": "全文",
        "headings_fixed": fixed_heading_count,
        "body_fixed": fixed_body_count,
        "target": {
            "body_font": f"宋体/TNR {body_size}pt",
            "line_spacing": f"{line_spacing}x",
            "first_indent": f"{first_indent_pt}pt",
        },
    })

    # ---------------------------------------------------------------
    # 3. Section titles (abstract, TOC, acknowledgments, appendices)
    # ---------------------------------------------------------------
    for keyword_group in ["abstract_zh", "abstract_en", "toc", "acknowledgments", "appendices"]:
        if keyword_group == "toc" and "toc" not in sections:
            continue
        if keyword_group == "acknowledgments" and "acknowledgments" not in sections:
            continue
        if keyword_group == "appendices" and "appendices" not in sections:
            continue
        if keyword_group.startswith("abstract") and "abstract" not in sections and "title_page" not in sections:
            continue

        for para in doc.paragraphs:
            if _is_section_title(para, keyword_group):
                old_size = _get_run_font_size(para)
                rec = _fix_section_title(para, old_size, keyword_group)
                if rec:
                    rec["category"] = "sections"
                    rec["item"] = f"{keyword_group}_title"
                    rec["location"] = para.text.strip()[:20]
                    records.append(rec)
                break  # Only fix first occurrence

    # ---------------------------------------------------------------
    # 4. Cover title fix (thesis mode only)
    # ---------------------------------------------------------------
    if "cover" in sections:
        COVER_KEYWORDS = ["学号", "姓名", "专业", "指导教师", "指导老师", "学院"]
        TITLE_KW = ["学士学位论文", "硕士学位论文", "博士学位论文", "毕业论文", "毕业设计"]

        cover_found = False
        for para in doc.paragraphs[:30]:
            text = para.text.strip()
            if not text:
                continue
            if any(kw in text for kw in COVER_KEYWORDS):
                cover_found = True
            if cover_found and any(kw in text for kw in TITLE_KW):
                old_size = _get_run_font_size(para)
                rec = _fix_section_title(para, old_size)
                if rec:
                    rec["category"] = "cover"
                    rec["item"] = "cover_title"
                    rec["location"] = text[:30]
                    records.append(rec)
                break

    # ---------------------------------------------------------------
    # 5. Citation repair
    # ---------------------------------------------------------------
    if repair_citations:
        try:
            ref_res = detect_references_section(str(working_path))
            if ref_res.found:
                cit_results = process_citations(ref_res.raw_citations)
                repair_citations_in_docx(str(working_path), cit_results, doc=doc)
                warns_count = sum(1 for r in cit_results if r.get("warnings"))
                records.append({
                    "category": "citations",
                    "item": "citation_repair",
                    "location": "参考文献",
                    "count": ref_res.citation_count,
                    "with_warnings": warns_count,
                    "results": [
                        {"original": r["original"][:80], "formatted": r.get("formatted", ""),
                         "warnings": r.get("warnings", [])}
                        for r in cit_results
                    ],
                })
        except Exception as e:
            records.append({"category": "citations", "item": "citation_repair",
                            "error": f"引用修复出错: {e}"})

    # ---------------------------------------------------------------
    # Reference font normalization (五号 10.5pt) — ALWAYS, independent of citation
    # reformatting. The references section must be 10.5pt regardless of whether
    # citation_repair ran, succeeded, or was skipped (e.g. repair_citations=False on
    # the .bib path). Doing it here means every path through fix_paper/fix_format
    # lands references at 10.5pt even if the (fragile) citation step bails out.
    # ---------------------------------------------------------------
    try:
        from scripts.citation_repair import _apply_citation_fonts, _is_list_item_stub
        ref_norm = detect_references_section(str(working_path))
        if ref_norm.found:
            start = ref_norm.section_start or 0
            end = ref_norm.section_end or len(doc.paragraphs)
            ref_count = 0
            for para in doc.paragraphs[start:end]:
                t = para.text.strip()
                if not t or _is_list_item_stub(t):
                    continue
                _apply_citation_fonts(para)  # 宋体+TNR 10.5pt
                ref_count += 1
            if ref_count:
                records.append({"category": "citations", "item": "reference_font",
                                "location": "参考文献", "note": f"统一字号为五号10.5pt（{ref_count} 条）"})
    except Exception as e:
        records.append({"category": "citations", "item": "reference_font", "error": str(e)})

    # ---------------------------------------------------------------
    # Save
    # ---------------------------------------------------------------
    try:
        doc.save(out_path)
    except PermissionError:
        return {"error": f"文件被占用或无写入权限: {out_path}", "output": None,
                "repair_records": records}
    except Exception as e:
        return {"error": f"保存失败: {e}", "output": None, "repair_records": records}

    return {
        "output": str(out_path),
        "mode": mode,
        "spec": str(spec_path) if spec_path else "default",
        "repair_records": records,
    }


# -------------------------------------------------------------------
# CLI
# -------------------------------------------------------------------

def main():
    from scripts.lib import setup_utf8_output
    setup_utf8_output()
    parser = argparse.ArgumentParser(description="Auto repair thesis format")
    parser.add_argument("thesis", help="thesis .docx path")
    parser.add_argument("--spec", "-s", help="Path to spec JSON (from parse_spec.py)")
    parser.add_argument("--output", "-o", help="output path")
    parser.add_argument("--no-citations", action="store_true", help="skip citation repair")
    parser.add_argument(
        "--mode", "-m", choices=["thesis", "journal"], default="journal",
        help="Document mode: thesis (default) or journal",
    )
    args = parser.parse_args()

    result = fix_format(
        args.thesis,
        spec_path=args.spec,
        output_path=args.output,
        repair_citations=not args.no_citations,
        mode=args.mode,
    )
    if result.get("error"):
        print(f"Error: {result['error']}")
    else:
        print(f"Output: {result['output']}")
        print(f"Records: {len(result['repair_records'])} repair actions")


if __name__ == "__main__":
    main()
