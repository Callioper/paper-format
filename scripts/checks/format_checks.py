"""format_checks.py - Individual format check functions.

Extracted from check_format.py for modularity.
Each function takes (doc, issues, spec) and populates the issues dict.
"""

from __future__ import annotations
import re
from typing import Any

# Default values (same as UNIVERSAL_CHECKS in check_format.py)
_DEFAULT_HEADING_SIZES = {1: 16, 2: 15, 3: 14, 4: 12}
_DEFAULT_BODY_SIZE = 12
_DEFAULT_LINE_SPACING = 1.5

def _pt(length) -> float | None:
    """Convert a python-docx Length to pt, or return None."""
    try:
        return length.pt if length is not None else None
    except Exception:
        return None


def _check_styles(doc, issues: dict, heading_sizes: dict | None = None) -> None:
    """Check Normal style font/size and Heading 1-4 font sizes."""
    from docx.oxml.ns import qn

    if heading_sizes is None:
        heading_sizes = _DEFAULT_HEADING_SIZES

    # --- Normal style body font ---
    try:
        normal = doc.styles["Normal"]
        nf = normal.font
        # CJK font via XML (python-docx doesn't expose eastAsia directly)
        rPr_elem = nf._element
        rFonts = rPr_elem.find(qn("w:rFonts")) if rPr_elem is not None else None
        east_asia = rFonts.get(qn("w:eastAsia")) if rFonts is not None else None
        ascii_f   = rFonts.get(qn("w:ascii"))    if rFonts is not None else None
        size_pt   = _pt(nf.size)

        if east_asia and east_asia not in ("宋体", "SimSun", "宋体,Bold"):
            issues["styles"].append({
                "item": "body_font_cjk",
                "expected": "宋体 (SimSun)",
                "actual": east_asia,
                "severity": "high",
                "suggestion": "正文 Normal 样式中文字体应设为宋体",
            })
        if ascii_f and "Times New Roman" not in ascii_f:
            issues["styles"].append({
                "item": "body_font_latin",
                "expected": "Times New Roman",
                "actual": ascii_f,
                "severity": "high",
                "suggestion": "正文 Normal 样式英文字体应设为 Times New Roman",
            })
        if size_pt is not None and abs(size_pt - _DEFAULT_BODY_SIZE) > 0.5:
            issues["styles"].append({
                "item": "body_font_size",
                "expected": f"{_DEFAULT_BODY_SIZE}pt (小四)",
                "actual": f"{size_pt:.1f}pt",
                "severity": "high",
                "suggestion": "正文字号应为 12pt（小四）",
            })
    except KeyError:
        pass  # Style not present

    # --- Normal style line spacing ---
    try:
        normal = doc.styles["Normal"]
        ls = normal.paragraph_format.line_spacing
        if ls is not None:
            from docx.enum.text import WD_LINE_SPACING
            rule = normal.paragraph_format.line_spacing_rule

            # Handle named rules first (python-docx special cases)
            if rule == WD_LINE_SPACING.ONE_POINT_FIVE:
                ratio = 1.5
            elif rule == WD_LINE_SPACING.DOUBLE:
                ratio = 2.0
            elif rule == WD_LINE_SPACING.AT_LEAST:
                # At-least: value is in EMU/pt, interpret as fixed
                ls_pt = _pt(ls)
                ratio = ls_pt / 12.0 if ls_pt else None
            elif rule == WD_LINE_SPACING.EXACTLY:
                ls_pt = _pt(ls)
                ratio = ls_pt / 12.0 if ls_pt else None
            elif rule == WD_LINE_SPACING.MULTIPLE:
                ratio = float(ls) if isinstance(ls, (int, float)) else None
            else:
                # Unknown rule, try to interpret value directly
                ratio = float(ls) if isinstance(ls, (int, float)) else None

            if ratio is not None and abs(ratio - _DEFAULT_LINE_SPACING) > 0.1:
                issues["styles"].append({
                    "item": "line_spacing",
                    "expected": f"{_DEFAULT_LINE_SPACING}x",
                    "actual": f"{ratio:.2f}x",
                    "severity": "high",
                    "suggestion": "行距应设为 1.5 倍",
                })
    except Exception:
        pass

    # --- Heading font sizes ---
    heading_sizes = dict(_DEFAULT_HEADING_SIZES)
    seen_headings: set[int] = set()
    for para in doc.paragraphs:
        sname = para.style.name if para.style else ""
        m = re.search(r"Heading\s*(\d)", sname, re.IGNORECASE)
        if not m:
            continue
        level = int(m.group(1))
        if level not in heading_sizes or level in seen_headings:
            continue
        seen_headings.add(level)
        expected_pt = heading_sizes[level]

        # Prefer run-level size, fall back to style size
        actual_pt: float | None = None
        for run in para.runs:
            if run.font.size:
                actual_pt = _pt(run.font.size)
                break
        if actual_pt is None:
            try:
                actual_pt = _pt(para.style.font.size)
            except Exception:
                pass

        if actual_pt is not None and abs(actual_pt - expected_pt) > 0.5:
            issues["styles"].append({
                "item": f"heading_{level}_size",
                "expected": f"{expected_pt}pt",
                "actual": f"{actual_pt:.1f}pt",
                "severity": "high",
                "suggestion": f"{'一二三四'[level-1]}级标题字号应为 {expected_pt}pt",
                "sample_text": para.text[:40],
            })

    # Warn for any required heading levels not found
    for level in heading_sizes:
        if level not in seen_headings:
            issues["styles"].append({
                "item": f"heading_{level}_missing",
                "expected": f"Heading {level} style used",
                "actual": "no paragraph found",
                "severity": "low",
                "suggestion": f"未找到{'一二三四'[level-1]}级标题段落，如文档有此层级请检查样式",
            })


def _check_body_paragraphs(doc, issues: dict) -> None:
    """
    Sample body paragraphs for first-line indent and line spacing overrides.
    Only report if violations are present in majority of sampled paragraphs.
    """
    from docx.shared import Pt as _Pt

    SAMPLE_LIMIT = 30
    indent_violations = 0
    spacing_violations = 0
    sampled = 0
    expected_indent_emu = int(_Pt(24))  # 2 × 12pt characters ≈ 24pt

    for para in doc.paragraphs:
        sname = para.style.name if para.style else ""
        if sname.startswith("Heading") or not para.text.strip():
            continue
        if len(para.text) < 10:  # skip short lines (captions, headers)
            continue
        sampled += 1
        if sampled > SAMPLE_LIMIT:
            break

        fi = para.paragraph_format.first_line_indent
        if fi is not None:
            fi_emu = int(fi)
            # Acceptable: 280000–380000 EMU (roughly 22–30pt, i.e. 1.8–2.4 chars)
            if fi_emu < 280000 or fi_emu > 400000:
                indent_violations += 1

        # Check paragraph-level line spacing overrides
        try:
            ls = para.paragraph_format.line_spacing
            if ls is not None:
                from docx.enum.text import WD_LINE_SPACING
                rule = para.paragraph_format.line_spacing_rule
                if rule == WD_LINE_SPACING.ONE_POINT_FIVE:
                    ratio = 1.5
                elif rule == WD_LINE_SPACING.DOUBLE:
                    ratio = 2.0
                elif rule == WD_LINE_SPACING.MULTIPLE:
                    ratio = float(ls)
                elif rule in (WD_LINE_SPACING.EXACTLY, WD_LINE_SPACING.AT_LEAST):
                    ratio = _pt(ls) / 12.0 if _pt(ls) else None
                else:
                    ratio = float(ls) if isinstance(ls, (int, float)) else None
                if ratio is not None and abs(ratio - _DEFAULT_LINE_SPACING) > 0.2:
                    spacing_violations += 1
        except Exception:
            pass

    if sampled > 5 and indent_violations / sampled > 0.4:
        issues["paragraphs"].append({
            "item": "first_line_indent",
            "expected": "2字符（约24pt）",
            "actual": f"{indent_violations}/{sampled} 段落不符",
            "severity": "medium",
            "suggestion": "正文段落首行缩进应为2字符（约0.74cm）",
        })

    if sampled > 5 and spacing_violations / sampled > 0.3:
        issues["paragraphs"].append({
            "item": "line_spacing_override",
            "expected": f"{_DEFAULT_LINE_SPACING}x",
            "actual": f"{spacing_violations}/{sampled} 段落有行距覆盖",
            "severity": "medium",
            "suggestion": "部分段落行距与 Normal 样式不一致，建议统一为 1.5 倍行距",
        })


def _check_cover(doc, issues: dict, spec: dict | None = None) -> None:
    """Check cover page format (thesis mode).

    Detects cover page by looking for typical keywords in the first 30 paragraphs.
    Checks title font/size/alignment and student info formatting.
    """
    from docx.shared import Pt

    COVER_KEYWORDS = ["学号", "姓名", "专业", "指导教师", "指导老师", "学院", "日期", "届别"]
    TITLE_KEYWORDS = ["学士学位论文", "硕士学位论文", "博士学位论文", "毕业论文", "毕业设计"]

    # Find cover page (first 30 paragraphs)
    cover_paras = []
    is_cover = False
    for i, para in enumerate(doc.paragraphs[:30]):
        text = para.text.strip()
        if not text:
            continue
        if any(kw in text for kw in COVER_KEYWORDS):
            is_cover = True
        cover_paras.append(para)

    if not is_cover:
        issues["cover"].append({
            "item": "cover_not_detected",
            "severity": "low",
            "suggestion": "未检测到封面页（未找到学号/姓名/专业等关键词），如论文有封面请检查",
        })
        return

    # Check cover title (usually the first non-empty centered paragraph)
    for para in cover_paras[:10]:
        text = para.text.strip()
        if not text or len(text) < 4:
            continue
        # Title is usually centered and contains thesis type keyword or is long enough
        is_title = any(kw in text for kw in TITLE_KEYWORDS) or len(text) > 15
        if not is_title:
            continue

        # Check alignment (should be centered)
        alignment = para.paragraph_format.alignment
        if alignment is not None and alignment != 1:  # 1 = CENTER
            issues["cover"].append({
                "item": "cover_title_alignment",
                "expected": "居中",
                "actual": "未居中",
                "severity": "medium",
                "suggestion": "封面论文标题应居中对齐",
                "sample_text": text[:30],
            })

        # Check font size
        for run in para.runs:
            if run.font.size:
                size_pt = _pt(run.font.size)
                if size_pt is not None and (size_pt < 14 or size_pt > 22):
                    issues["cover"].append({
                        "item": "cover_title_size",
                        "expected": "小二(18pt) 或 二号(22pt)",
                        "actual": f"{size_pt:.0f}pt",
                        "severity": "medium",
                        "suggestion": "封面标题字号通常为小二或二号",
                        "sample_text": text[:30],
                    })
                break
        break  # Only check the first title-like paragraph


def _check_abstract(doc, issues: dict, spec: dict | None = None) -> None:
    """Check abstract and keywords format."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ABSTRACT_KEYWORDS = ["摘  要", "摘 要", "摘  要", "摘  要", "Abstract", "ABSTRACT"]
    KEYWORD_LABELS = ["关键词", "Keywords", "Key words", "KEYWORDS"]

    found_abstract_title = False
    in_abstract = False
    found_keywords = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Detect abstract title
        if not found_abstract_title and any(kw in text for kw in ABSTRACT_KEYWORDS):
            found_abstract_title = True
            in_abstract = True

            # Check abstract title format
            alignment = para.paragraph_format.alignment
            if alignment is not None and alignment != 1:
                issues["abstract"].append({
                    "item": "abstract_title_alignment",
                    "expected": "居中",
                    "actual": "未居中",
                    "severity": "medium",
                    "suggestion": "摘要标题应居中对齐",
                    "sample_text": text[:20],
                })

            for run in para.runs:
                if run.font.size:
                    size_pt = _pt(run.font.size)
                    if size_pt is not None and abs(size_pt - 16) > 1:
                        issues["abstract"].append({
                            "item": "abstract_title_size",
                            "expected": "三号(16pt)",
                            "actual": f"{size_pt:.0f}pt",
                            "severity": "medium",
                            "suggestion": "摘要标题字号应为三号(16pt)",
                            "sample_text": text[:20],
                        })
                    break
            continue

        # Detect keywords line
        if in_abstract and any(kw in text for kw in KEYWORD_LABELS):
            found_keywords = True
            in_abstract = False

            # Check keyword label is bold
            label_bold = False
            for run in para.runs:
                if run.font.bold:
                    label_bold = True
                    break
            if not label_bold:
                issues["abstract"].append({
                    "item": "keywords_label_bold",
                    "expected": "关键词标签加粗",
                    "actual": "未加粗",
                    "severity": "low",
                    "suggestion": "「关键词」标签应加粗",
                })

            # Check separator (should use semicolons or Chinese semicolons)
            kw_text = text.split("：" if "：" in text else ":")[-1] if ("：" in text or ":" in text) else ""
            if kw_text and "；" not in kw_text and ";" not in kw_text:
                issues["abstract"].append({
                    "item": "keywords_separator",
                    "expected": "关键词间用分号分隔",
                    "actual": "未检测到分号",
                    "severity": "low",
                    "suggestion": "多个关键词之间应用分号（；）分隔",
                })
            continue

    if not found_abstract_title:
        issues["abstract"].append({
            "item": "abstract_not_found",
            "severity": "medium",
            "suggestion": "未检测到摘要章节，请检查是否存在「摘 要」或「Abstract」标题",
        })


def _check_toc(doc, issues: dict, spec: dict | None = None) -> None:
    """Check table of contents format (thesis mode)."""
    TOC_KEYWORDS = ["目  录", "目 录", "目    录", "Contents", "CONTENTS"]

    found_toc = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if any(kw in text for kw in TOC_KEYWORDS):
            found_toc = True

            # Check alignment
            alignment = para.paragraph_format.alignment
            if alignment is not None and alignment != 1:
                issues["toc"].append({
                    "item": "toc_title_alignment",
                    "expected": "居中",
                    "actual": "未居中",
                    "severity": "medium",
                    "suggestion": "目录标题应居中对齐",
                    "sample_text": text[:20],
                })

            # Check font size
            for run in para.runs:
                if run.font.size:
                    size_pt = _pt(run.font.size)
                    if size_pt is not None and abs(size_pt - 16) > 1:
                        issues["toc"].append({
                            "item": "toc_title_size",
                            "expected": "三号(16pt)",
                            "actual": f"{size_pt:.0f}pt",
                            "severity": "medium",
                            "suggestion": "目录标题字号应为三号(16pt)",
                            "sample_text": text[:20],
                        })
                    break
            break

    if not found_toc:
        issues["toc"].append({
            "item": "toc_not_found",
            "severity": "low",
            "suggestion": "未检测到目录页（未找到「目 录」或「Contents」标题）",
        })


def _check_tables(doc, issues: dict, spec: dict | None = None) -> None:
    """Check table border styles (three-line table)."""
    from docx.oxml.ns import qn

    for idx, table in enumerate(doc.tables):
        tbl_pr = table._tbl.tblPr
        if tbl_pr is None:
            continue

        borders_el = tbl_pr.find(qn("w:tblBorders"))
        if borders_el is None:
            issues["tables"].append({
                "item": f"table_{idx}_no_borders",
                "severity": "low",
                "suggestion": f"表格{idx+1}未设置边框样式",
            })
            continue

        border_vals = {}
        for side in ("top", "bottom", "left", "right", "insideH", "insideV"):
            side_el = borders_el.find(qn(f"w:{side}"))
            if side_el is not None:
                val = side_el.get(qn("w:val"), "none")
                sz = side_el.get(qn("w:sz"), "0")
                border_vals[side] = {"val": val, "sz": int(sz) if sz.isdigit() else 0}
            else:
                border_vals[side] = {"val": "none", "sz": 0}

        has_top = border_vals["top"]["val"] not in ("none", "nil")
        has_bottom = border_vals["bottom"]["val"] not in ("none", "nil")
        no_left = border_vals["left"]["val"] in ("none", "nil")
        no_right = border_vals["right"]["val"] in ("none", "nil")

        if not (has_top and has_bottom and no_left and no_right):
            issues["tables"].append({
                "item": f"table_{idx}_style",
                "expected": "三线表（仅顶线+底线+表头底线）",
                "actual": f"顶线={'有' if has_top else '无'}, 底线={'有' if has_bottom else '无'}, "
                          f"左线={'有' if not no_left else '无'}, 右线={'有' if not no_right else '无'}",
                "severity": "medium",
                "suggestion": f"表格{idx+1}不是三线表格式，建议使用三线表",
            })


def _check_headers_footers(doc, issues: dict, spec: dict | None = None) -> None:
    """Check header and footer format (thesis mode)."""
    from docx.oxml.ns import qn

    for idx, section in enumerate(doc.sections):
        try:
            header = section.header
            if header and header.paragraphs:
                for para in header.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    # Check font size (should be 五号 = 10.5pt)
                    for run in para.runs:
                        if run.font.size:
                            size_pt = _pt(run.font.size)
                            if size_pt is not None and abs(size_pt - 10.5) > 1:
                                issues["headers_footers"].append({
                                    "item": f"header_{idx}_font_size",
                                    "expected": "五号(10.5pt)",
                                    "actual": f"{size_pt:.1f}pt",
                                    "severity": "low",
                                    "suggestion": "页眉字号应为五号(10.5pt)",
                                    "sample_text": text[:30],
                                })
                            break
                    break  # Only check first non-empty paragraph

            # Check for header bottom border (页眉横线)
            for para in header.paragraphs:
                pPr = para._p.get_or_add_pPr()
                pBdr = pPr.find(qn("w:pBdr"))
                if pBdr is None:
                    # No border on header — this is common, report as info
                    pass
        except Exception:
            pass


def _check_acknowledgments(doc, issues: dict, spec: dict | None = None) -> None:
    """Check acknowledgments section format (thesis mode)."""
    ACK_KEYWORDS = ["致  谢", "致 谢", "致    谢"]

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if any(kw in text for kw in ACK_KEYWORDS):
            # Check alignment
            alignment = para.paragraph_format.alignment
            if alignment is not None and alignment != 1:
                issues["acknowledgments"].append({
                    "item": "ack_title_alignment",
                    "expected": "居中",
                    "actual": "未居中",
                    "severity": "medium",
                    "suggestion": "致谢标题应居中对齐",
                    "sample_text": text[:20],
                })

            # Check font size
            for run in para.runs:
                if run.font.size:
                    size_pt = _pt(run.font.size)
                    if size_pt is not None and abs(size_pt - 16) > 1:
                        issues["acknowledgments"].append({
                            "item": "ack_title_size",
                            "expected": "三号(16pt)",
                            "actual": f"{size_pt:.0f}pt",
                            "severity": "medium",
                            "suggestion": "致谢标题字号应为三号(16pt)",
                            "sample_text": text[:20],
                        })
                    break
            break


def _check_appendices(doc, issues: dict, spec: dict | None = None) -> None:
    """Check appendices section format (thesis mode)."""
    APP_KEYWORDS = ["附  录", "附 录", "附    录", "Appendix", "APPENDIX"]

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if any(kw in text for kw in APP_KEYWORDS):
            # Check alignment
            alignment = para.paragraph_format.alignment
            if alignment is not None and alignment != 1:
                issues["appendices"].append({
                    "item": "appendix_title_alignment",
                    "expected": "居中",
                    "actual": "未居中",
                    "severity": "medium",
                    "suggestion": "附录标题应居中对齐",
                    "sample_text": text[:20],
                })

            # Check font size
            for run in para.runs:
                if run.font.size:
                    size_pt = _pt(run.font.size)
                    if size_pt is not None and abs(size_pt - 16) > 1:
                        issues["appendices"].append({
                            "item": "appendix_title_size",
                            "expected": "三号(16pt)",
                            "actual": f"{size_pt:.0f}pt",
                            "severity": "medium",
                            "suggestion": "附录标题字号应为三号(16pt)",
                            "sample_text": text[:20],
                        })
                    break
            break


# -------------------------------------------------------------------
# Main check function
# -------------------------------------------------------------------
