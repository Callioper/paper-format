#!/usr/bin/env python3
"""check_content.py — 中文学术论文内容结构检查。

检查：必需章节是否存在、摘要关键词数量、缩略语候选识别。

Usage:
    python scripts/check_content.py "论文.docx" --mode journal --output content_result.json
    python scripts/check_content.py "论文.docx" --mode thesis
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

try:
    from docx import Document
except ImportError:
    print("需要 python-docx: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ── 必需章节定义 ──────────────────────────────────────────

_JOURNAL_SECTIONS = {
    "required": [
        {"name": "标题", "patterns": [r"^.{2,50}$"], "check": "first_paragraph"},
        {"name": "摘要", "patterns": [r"摘\s*要", "摘　要", "Abstract", "ABSTRACT"]},
        {"name": "关键词", "patterns": [r"关\s*键\s*词", "关键词", "Keywords", "Key words"]},
        {"name": "正文", "patterns": ["正文", "引言", "绪论", r"1\s", "一、"]},
        {"name": "参考文献", "patterns": ["参考文献", "References", "引用", "Works Cited", "Bibliography"]},
    ],
    "optional": [
        {"name": "英文摘要", "patterns": ["Abstract", "ABSTRACT"]},
        {"name": "作者信息", "patterns": ["作者", "Author", "作者简介"]},
        {"name": "基金项目", "patterns": ["基金", "Fund", "项目编号", "基金项目"]},
    ],
}

_THESIS_SECTIONS = {
    "required": [
        {"name": "封面", "patterns": ["封面", "Cover", "题名页"], "check": "first_pages"},
        {"name": "中文摘要", "patterns": [r"摘\s*要", "摘　要"]},
        {"name": "英文摘要", "patterns": ["Abstract", "ABSTRACT"]},
        {"name": "关键词", "patterns": [r"关\s*键\s*词", "Keywords"]},
        {"name": "目录", "patterns": [r"目\s*录", "目　录", "Table of Contents"]},
        {"name": "正文", "patterns": [r"第[一二三四五六七八九十1-9]", "Chapter", "绪论", "引言"]},
        {"name": "参考文献", "patterns": ["参考文献", "References", "Bibliography"]},
        {"name": "致谢", "patterns": [r"致\s*谢", "致　谢", "Acknowledgment"]},
    ],
    "optional": [
        {"name": "附录", "patterns": [r"附\s*录", "Appendix"]},
        {"name": "作者简介", "patterns": ["作者简介", "Author Bio"]},
        {"name": "声明", "patterns": [r"声\s*明", "Declaration"]},
    ],
}


# ── 检查函数 ──────────────────────────────────────────────

def _extract_headings(doc: Document) -> list[dict[str, Any]]:
    """提取所有标题段落。"""
    headings = []
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ""
        if "Heading" in style_name or "标题" in style_name:
            headings.append({
                "index": i,
                "text": para.text.strip(),
                "style": style_name,
            })
    return headings


def _check_section_exists(headings: list[dict], section_def: dict, all_text: str) -> dict[str, Any]:
    """检查某个章节是否存在。"""
    patterns = section_def["patterns"]
    found = False
    match_info = ""

    # 先在标题中搜索
    for h in headings:
        for pat in patterns:
            if re.search(pat, h["text"], re.IGNORECASE):
                found = True
                match_info = f"标题: '{h['text']}'"
                break
        if found:
            break

    # 如果标题中没找到，在全文中搜索
    if not found:
        for pat in patterns:
            m = re.search(pat, all_text, re.IGNORECASE)
            if m:
                found = True
                match_info = f"全文位置 {m.start()}"
                break

    return {
        "name": section_def["name"],
        "found": found,
        "match": match_info,
        "required": "required" in str(section_def.get("_type", "")),
    }


def _count_keywords(all_text: str) -> dict[str, Any]:
    """检测摘要关键词数量。"""
    # 查找关键词行
    kw_match = re.search(
        r"(?:关\s*键\s*词|关键词|Keywords|Key\s*[Ww]ords)[：:\s]*(.+?)(?:\n|$)",
        all_text, re.IGNORECASE
    )
    if not kw_match:
        return {"found": False, "count": 0, "keywords": [], "recommendation": "未找到关键词行"}

    kw_text = kw_match.group(1).strip()
    # 按常见分隔符拆分
    keywords = re.split(r"[;；,，、\s]+", kw_text)
    keywords = [k.strip() for k in keywords if k.strip() and len(k.strip()) > 1]

    recommendation = ""
    count = len(keywords)
    if count < 3:
        recommendation = "关键词数量偏少（建议 3-8 个）"
    elif count > 8:
        recommendation = "关键词数量偏多（建议 3-8 个）"
    else:
        recommendation = "关键词数量合适"

    return {
        "found": True,
        "count": count,
        "keywords": keywords,
        "recommendation": recommendation,
    }


def _find_acronyms(all_text: str) -> list[dict[str, str]]:
    """识别可能的缩略语候选。"""
    acronyms = []
    # 匹配全大写缩略语（2-6 个字母）
    pattern = r"\b([A-Z]{2,6})\b"
    seen = set()
    for m in re.finditer(pattern, all_text):
        abbr = m.group(1)
        if abbr in seen:
            continue
        # 排除常见非缩略语
        if abbr in {"THE", "AND", "FOR", "BUT", "NOT", "ARE", "WAS", "HAS", "HAD",
                     "ITS", "OUR", "ALL", "CAN", "MAY", "WILL", "SHALL", "THIS",
                     "THAT", "WITH", "FROM", "HAVE", "BEEN", "WERE", "THEY",
                     "WHAT", "WHEN", "WHERE", "HOW", "WHO", "WHICH"}:
            continue
        seen.add(abbr)
        # 查找上下文中的可能全称
        context = all_text[max(0, m.start() - 50):m.end() + 50]
        acronyms.append({
            "abbreviation": abbr,
            "position": m.start(),
            "context": context[:60],
        })

    return acronyms


def _count_abstract_words(all_text: str) -> dict[str, Any]:
    """检测摘要字数。"""
    # 查找摘要段落
    abs_match = re.search(
        r"(?:摘\s*要|Abstract)[：:\s]*(.+?)(?=(?:关\s*键\s*词|Keywords|Key\s*[Ww]ords|正文|引言|\n\n\n))",
        all_text, re.IGNORECASE | re.DOTALL
    )
    if not abs_match:
        return {"found": False, "char_count": 0, "recommendation": "未找到摘要"}

    abstract_text = abs_match.group(1).strip()
    char_count = len(abstract_text)

    recommendation = ""
    if char_count < 100:
        recommendation = "摘要过短（建议中文摘要 200-300 字）"
    elif char_count > 500:
        recommendation = "摘要过长（建议中文摘要 200-300 字）"
    else:
        recommendation = "摘要长度合适"

    return {
        "found": True,
        "char_count": char_count,
        "recommendation": recommendation,
    }


# ── 主检查流程 ──────────────────────────────────────────────

def check_content(docx_path: str | Path, mode: str = "journal") -> dict[str, Any]:
    """执行内容结构检查。"""
    doc = Document(str(docx_path))
    headings = _extract_headings(doc)

    # 提取全文文本
    all_text = "\n".join(p.text for p in doc.paragraphs)

    # 选择章节定义
    section_defs = _THESIS_SECTIONS if mode == "thesis" else _JOURNAL_SECTIONS

    # 检查必需章节
    required_results = []
    for sec in section_defs["required"]:
        sec["_type"] = "required"
        result = _check_section_exists(headings, sec, all_text)
        result["required"] = True
        required_results.append(result)

    # 检查可选章节
    optional_results = []
    for sec in section_defs["optional"]:
        sec["_type"] = "optional"
        result = _check_section_exists(headings, sec, all_text)
        result["required"] = False
        optional_results.append(result)

    # 关键词检查
    keywords_info = _count_keywords(all_text)

    # 摘要字数检查
    abstract_info = _count_abstract_words(all_text)

    # 缩略语候选
    acronyms = _find_acronyms(all_text)

    # 汇总
    missing_required = [r["name"] for r in required_results if not r["found"]]
    total_issues = len(missing_required)
    if keywords_info.get("count", 0) < 3:
        total_issues += 1
    if abstract_info.get("char_count", 0) < 100:
        total_issues += 1

    return {
        "file": str(docx_path),
        "mode": mode,
        "total_issues": total_issues,
        "required_sections": required_results,
        "optional_sections": optional_results,
        "missing_required": missing_required,
        "keywords": keywords_info,
        "abstract": abstract_info,
        "acronym_candidates": acronyms[:20],  # 最多报告 20 个
        "heading_count": len(headings),
        "headings": [h["text"] for h in headings[:30]],  # 前 30 个标题
    }


# ── CLI ──────────────────────────────────────────────────

def main() -> int:
    parser = argparse.ArgumentParser(description="中文学术论文内容结构检查")
    parser.add_argument("docx", help="待检查的 .docx 文件")
    parser.add_argument("--mode", "-m", default="journal", choices=["journal", "thesis"],
                        help="文档模式：journal（期刊论文）或 thesis（学位论文）")
    parser.add_argument("--output", "-o", help="输出 JSON 路径")
    args = parser.parse_args()

    docx_path = Path(args.docx)
    if not docx_path.exists():
        print(f"文件不存在: {docx_path}", file=sys.stderr)
        return 1

    report = check_content(docx_path, args.mode)

    if args.output:
        Path(args.output).write_text(
            json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        print(f"检查完成: {report['total_issues']} 个问题 → {args.output}")
    else:
        print(json.dumps(report, ensure_ascii=False, indent=2))

    # 打印摘要
    print(f"\n=== 内容结构检查摘要 ({args.mode}) ===")
    for sec in report["required_sections"]:
        status = "[OK]" if sec["found"] else "[MISSING]"
        print(f"  {status} {sec['name']}: {sec.get('match', 'not found')}")
    if report["missing_required"]:
        print(f"\n  缺失必需章节: {', '.join(report['missing_required'])}")
    kw = report["keywords"]
    if kw["found"]:
        print(f"\n  关键词 ({kw['count']} 个): {', '.join(kw['keywords'])}")
        print(f"  {kw['recommendation']}")
    abs_info = report["abstract"]
    if abs_info["found"]:
        print(f"\n  摘要字数: {abs_info['char_count']}")
        print(f"  {abs_info['recommendation']}")
    if report["acronym_candidates"]:
        print(f"\n  缩略语候选 ({len(report['acronym_candidates'])} 个): "
              f"{', '.join(a['abbreviation'] for a in report['acronym_candidates'][:10])}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
