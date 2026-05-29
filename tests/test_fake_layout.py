import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from scripts.check_fake_layout import detect_fake_layout, clean_fake_layout


def _doc_to_tmp(doc, tmp_path):
    p = tmp_path / "t.docx"
    doc.save(str(p))
    return str(p)


def test_detects_consecutive_blank_paragraphs(tmp_path):
    doc = Document()
    doc.add_paragraph("正文第一段。")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")          # 3 连续空段
    doc.add_paragraph("正文第二段。")
    res = detect_fake_layout(_doc_to_tmp(doc, tmp_path))
    blanks = [i for i in res["issues"] if i["type"] == "consecutive_blank"]
    assert len(blanks) == 1
    assert blanks[0]["count"] == 3


def test_detects_leading_manual_spaces(tmp_path):
    doc = Document()
    doc.add_paragraph("　　这段用了两个全角空格冒充首行缩进。")  # 全角空格
    doc.add_paragraph("    这段用了四个半角空格。")
    res = detect_fake_layout(_doc_to_tmp(doc, tmp_path))
    leads = [i for i in res["issues"] if i["type"] == "manual_indent_spaces"]
    assert len(leads) == 2


def test_detects_intra_paragraph_break(tmp_path):
    doc = Document()
    p = doc.add_paragraph("第一行")
    br = OxmlElement("w:br")
    p.runs[0]._r.append(br)
    p.add_run("被手工换行挤到第二行")
    res = detect_fake_layout(_doc_to_tmp(doc, tmp_path))
    assert any(i["type"] == "intra_para_break" for i in res["issues"])


def test_detects_trailing_space(tmp_path):
    doc = Document()
    doc.add_paragraph("这段行尾有游离空格。   ")
    res = detect_fake_layout(_doc_to_tmp(doc, tmp_path))
    assert any(i["type"] == "trailing_space" for i in res["issues"])


def test_clean_doc_has_no_issues(tmp_path):
    doc = Document()
    doc.add_paragraph("正常段落一。")
    doc.add_paragraph("正常段落二。")
    res = detect_fake_layout(_doc_to_tmp(doc, tmp_path))
    assert res["total"] == 0


def test_clean_collapses_blanks_and_strips_leading_spaces(tmp_path):
    doc = Document()
    doc.add_paragraph("正文第一段。")
    doc.add_paragraph(""); doc.add_paragraph(""); doc.add_paragraph("")
    doc.add_paragraph("　　第二段带全角空格。")
    n = clean_fake_layout(doc)
    out = tmp_path / "cleaned.docx"; doc.save(str(out))
    d2 = Document(str(out))
    texts = [p.text for p in d2.paragraphs]
    # 连续空段被折叠为至多 1 个
    assert "\n".join(texts).count("\n\n\n") == 0
    # 段首全角空格被去除
    assert any(t.startswith("第二段") for t in texts)
    assert n >= 2  # 至少清理了 空段折叠 + 段首空格
